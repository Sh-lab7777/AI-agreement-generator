"""
pustaka_app.py — Pustaka Agreement Generator
Standalone CustomTkinter desktop app. No browser. No Streamlit.
Bundle with PyInstaller → single .exe for office staff.
"""

import os, sys, re, io, json, copy, threading, uuid
from lxml import etree as ET
import tkinter as tk
import tkinter.filedialog as fd
import tkinter.messagebox as mb
import customtkinter as ctk
from datetime import datetime
from docx import Document
import PyPDF2
import anthropic
import platform

def _config_dir():
    """Per-user config directory for storing API key."""
    if platform.system() == "Windows":
        base = os.environ.get("APPDATA", os.path.expanduser("~"))
    elif platform.system() == "Darwin":
        base = os.path.expanduser("~/Library/Application Support")
    else:
        base = os.path.expanduser("~/.config")
    d = os.path.join(base, "PustakaApp")
    os.makedirs(d, exist_ok=True)
    return d

def _key_file():
    return os.path.join(_config_dir(), "api_key.txt")

def _load_api_key():
    """Load API key: env var > .env file > saved config file > hardcoded."""
    # 1. Already set in environment
    if os.environ.get("ANTHROPIC_API_KEY","").strip():
        return os.environ["ANTHROPIC_API_KEY"].strip()
    # 2. .env file next to exe / script
    try:
        from dotenv import load_dotenv
        exe_dir = os.path.dirname(os.path.abspath(sys.executable if getattr(sys,"frozen",False) else __file__))
        load_dotenv(os.path.join(exe_dir, ".env"))
        if os.environ.get("ANTHROPIC_API_KEY","").strip():
            return os.environ["ANTHROPIC_API_KEY"].strip()
    except Exception:
        pass
    # 3. Saved config file
    try:
        with open(_key_file(), "r") as f:
            key = f.read().strip()
            if key:
                os.environ["ANTHROPIC_API_KEY"] = key
                return key
    except Exception:
        pass
    # 4. No hardcoded key — return empty, user must provide via .env or settings
    return ""

def _save_api_key(key):
    key = key.strip()
    with open(_key_file(), "w") as f:
        f.write(key)
    os.environ["ANTHROPIC_API_KEY"] = key

# Load key at startup
_load_api_key()

# ── PyInstaller resource path ──────────────────────────────────
def _res(rel):
    base = getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base, rel)

# ── Theme ──────────────────────────────────────────────────────
ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")

BLUE_DEEP   = "#0F2D6B"
BLUE_MID    = "#1A4FAA"
BLUE_LIGHT  = "#3B7DD8"
BLUE_PALE   = "#D6E6FF"
MINT        = "#D4F5E9"
BG          = "#F0F4FF"
WHITE       = "#FFFFFF"
TEXT_MUTED  = "#5B7AA8"
BORDER      = "#C5D8FF"
SUCCESS     = "#0EA47A"
WARNING     = "#F59E0B"
ERROR_RED   = "#DC2626"

FONT_TITLE  = ("Segoe UI", 20, "bold")
FONT_HEAD   = ("Segoe UI", 13, "bold")
FONT_BODY   = ("Segoe UI", 12)
FONT_SMALL  = ("Segoe UI", 10)
FONT_MONO   = ("Consolas", 11)

# ── Template paths ─────────────────────────────────────────────
TDIR = _res("templates")

KATHA_TEMPLATES = {
    "silver++": os.path.join(TDIR, "Silver ++ template - katha agreement.docx"),
    "pearl":    os.path.join(TDIR, "Silver ++ template - katha agreement.docx"),
    "sapphire": os.path.join(TDIR, "Silver ++ template - katha agreement.docx"),
    "sgp":      os.path.join(TDIR, "silver gold plan template - katha agreement.docx"),
    "silver":   os.path.join(TDIR, "silver gold plan template - katha agreement.docx"),
    "gold":     os.path.join(TDIR, "silver gold plan template - katha agreement.docx"),
}
PUSTAKA_TEMPLATES = {
    "new_author": os.path.join(TDIR, "Template_New_Authors_Paperback - pustaka agreement.docx"),
    "legal_heir": os.path.join(TDIR, "Template_Legal Heir_Digital - pustaka agreement.docx"),
    "licensor":   os.path.join(TDIR, "Template_Licensor_Paperback new - pustaka agreement.docx"),
}
ADDENDUM_TEMPLATES = {
    "katha_pearl":    os.path.join(TDIR, "pearl plan template (addendum).docx"),
    "katha_sapphire": os.path.join(TDIR, "sapphire plan template (addendum).docx"),
    "pustaka":        os.path.join(TDIR, "pustaka- defualt addendum.docx"),
}

# ── Field definitions ──────────────────────────────────────────
COMMON_FIELDS = {
    "author_name":          ("Author Name",                  "e.g. Suresh Kumar",               True),
    "pen_name":             ("Pen Name (if any)",           "e.g. Raja (optional)",             False),
    "parent_name":          ("Father / Husband Name",       "e.g. Mr. Komal Raj",              True),
    "address":              ("Full Postal Address",         "e.g. 45/C, Anna Nagar, Chennai",  True),
    "mobile":               ("Mobile No",                   "e.g. 9865230147",                 True),
    "email":                ("Email Id",                    "e.g. author@email.com",           True),
    "nominee_name":         ("Nominee Name",                "e.g. Priya Suresh",               True),
    "nominee_relationship": ("Nominee Relationship",        "e.g. Daughter / Wife / Son",      True),
    "nominee_address":      ("Nominee Full Postal Address", "e.g. Same as author address",     True),
    "nominee_age":          ("Nominee Age",                 "e.g. 28",                         True),
    "nominee_mobile_email": ("Nominee Mobile / Email",      "e.g. 9876543210 / priya@mail.com",True),
    "nationality":          ("Nationality",                 "e.g. Indian",                     True),
}
BOOK_FIELDS = {
    # book titles and genres are now collected via Annexure A checkboxes
}
KATHA_EXTRA = {
    "ref_number":   ("Reference Number",  "e.g. KAT/TAM/029/2026",          True),
    "date":         ("Agreement Date",    "e.g. 23rd February 2026",        True),
    "language":     ("Language",          "e.g. Tamil",                     True),
    "plan":         ("Publishing Plan",   "Pearl / Sapphire / Silver++ / SGP",True),
    "format":       ("Publish Format",    "e.g. Ebook and Paperback",       True),
    "amount":       ("Amount",            "e.g. Rs.7999/-",                 True),
    "sending_date": ("Sending Date",      "e.g. 26 March 2026",             True),
    "return_date":  ("Return Date",       "e.g. 02 April 2026",             True),
}
PUSTAKA_EXTRA = {
    "ref_number":   ("Reference Number",  "e.g. PUS/TAM/123/2026",  True),
    "date":         ("Agreement Date",    "e.g. 26th March 2026",   True),
    "language":     ("Language",          "e.g. Tamil",              True),
    "sending_date": ("Sending Date",      "e.g. 26 March 2026",     True),
    "return_date":  ("Return Date",       "e.g. 02 April 2026",     True),
}
ADDENDUM_FIELDS = {
    "ref_number":          ("Addendum Ref Number",      "e.g. KAT/TAM/325-04/2026",  True),
    "addendum_date":       ("Addendum Date",            "e.g. 23rd February 2026",   True),
    "orig_agreement_ref":  ("Original Agreement Ref",   "e.g. KAT/TAM/325-03/2026",  True),
    "orig_agreement_date": ("Original Agreement Date",  "e.g. 24th January 2026",    True),
    "author_name":         ("Author Full Name",         "e.g. Mr. C. Arockiya Samy", True),
    "pen_name":            ("Pen Name",                 "e.g. Vengai Aron (optional)",False),
    "parent_name":         ("Father/Mother Name",       "e.g. Mr. Cinnappan",         True),
    "address":             ("Address",                  "e.g. 1/120, Venkitakulam",   True),
    "email":               ("Email",                    "e.g. author@email.com",       True),
    "mobile":              ("Mobile Number(s)",         "e.g. 9786505016",             True),
    "language":            ("Language",                  "e.g. Tamil",                 True),
    "plan":                ("Publishing Plan",           "e.g. Sapphire / Pearl",      True),
    "amount":              ("Amount",                    "e.g. Rs.6999/-",             True),
    "format":              ("Publish Format",            "e.g. Ebook & Paperback",     True),
    "sending_date":        ("Sending Date",              "e.g. 26 March 2026",         True),
    "return_date":         ("Return Date",               "e.g. 02 April 2026",         True),
}

def get_fields(doc_type):
    if doc_type == "katha_agreement":
        return {**COMMON_FIELDS, **BOOK_FIELDS, **KATHA_EXTRA}
    elif doc_type == "pustaka_agreement":
        return {**COMMON_FIELDS, **PUSTAKA_EXTRA}
    elif doc_type == "addendum":
        return ADDENDUM_FIELDS
    return {}

# ── Doc generation helpers ─────────────────────────────────────
def merge_runs(para):
    if not para.runs: return
    full = "".join(r.text for r in para.runs)
    para.runs[0].text = full
    for r in para.runs[1:]: r.text = ""

def set_text(para, text):
    if not para.runs: para.add_run(text)
    else:
        para.runs[0].text = text
        for r in para.runs[1:]: r.text = ""

def wc(cell, text):
    p = cell.paragraphs[0]
    if not p.runs: p.add_run(text)
    else:
        p.runs[0].text = text
        for r in p.runs[1:]: r.text = ""

def get_book_list(fields):
    books = []
    for i in range(1, 4):
        val = fields.get(f"book_title_{i}", "").strip()
        if val:
            for sep in ["–", "-", "—"]:
                if sep in val:
                    parts = val.split(sep, 1)
                    books.append((parts[0].strip(), parts[1].strip()))
                    break
            else:
                books.append((val, ""))
    return books

def build_format_str(fmts):
    if not fmts: return "Ebook"
    m = {"ebook": "Ebook", "audiobook": "Audiobook", "paperback": "Paperback"}
    return " & ".join(m.get(f, f) for f in fmts)

FMT_LABEL = {"ebook": "Ebook", "audiobook": "Audiobook", "paperback": "Paperback"}
FMT_PARA  = {"ebook": "For Ebook Title", "audiobook": "For Audiobook Title", "paperback": "For Paperback Title"}

def _para_text(para):
    return "".join(r.text for r in para.runs).strip()

def _get_preceding_para_text(doc_element, tbl_element):
    """Return the text of the paragraph immediately before a table element."""
    from docx.oxml.ns import qn as _qn
    body = doc_element.body
    children = list(body)
    idx = children.index(tbl_element)
    for back in range(1, 6):
        if idx - back < 0: break
        prev = children[idx - back]
        if prev.tag.endswith("}p"):
            texts = [r.text for r in prev.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")]
            full = " ".join(t for t in texts if t).strip()
            if full:
                return full
    return ""

def _remove_element(element):
    """Remove an XML element from its parent."""
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)

def _remove_preceding_heading_paras(doc_element, tbl_element):
    """Remove 'For Ebook Title' / 'For Paperback Title' etc heading paragraphs before a table."""
    from docx.oxml.ns import qn as _qn
    body = doc_element.body
    children = list(body)
    idx = children.index(tbl_element)
    to_remove = []
    for back in range(1, 4):
        if idx - back < 0: break
        prev = children[idx - back]
        tag = prev.tag.split("}")[-1]
        if tag == "p":
            texts = [r.text for r in prev.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")]
            full = " ".join(t for t in texts if t).strip()
            if full.startswith("For ") and "Title" in full:
                to_remove.append(prev)
            elif not full:
                to_remove.append(prev)
            else:
                break
        else:
            break
    for el in to_remove:
        _remove_element(el)

def fill_annexure_a(doc, annexure_data, lang):
    """
    Fill Annexure A tables in doc.
    annexure_data = { 'ebook': [{title,language,genre,format,royalty},...], 'paperback': [...], 'audiobook': [...] }
    Only selected formats get filled; others are removed along with their heading paragraphs.
    """
    from copy import deepcopy
    from docx.oxml.ns import qn
    from lxml import etree

    FORMAT_HEADING_KEYWORDS = {
        "ebook":     "ebook",
        "paperback": "paperback",
        "audiobook": "audiobook",
    }

    body = doc.element.body
    # Find Annexure A tables (those preceded by "For ... Title" paragraphs)
    annexure_tables = {}  # fmt_key -> table_element
    for tbl_el in body.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl"):
        prev_text = _get_preceding_para_text(doc.element, tbl_el).lower()
        for fmt_key, kw in FORMAT_HEADING_KEYWORDS.items():
            if kw in prev_text and "title" in prev_text:
                annexure_tables[fmt_key] = tbl_el
                break

    for fmt_key, tbl_el in annexure_tables.items():
        rows = annexure_data.get(fmt_key, [])
        has_data = bool(rows and any(r.get("title","").strip() for r in rows))

        if not has_data:
            # Remove this table and its heading para
            _remove_preceding_heading_paras(doc.element, tbl_el)
            _remove_element(tbl_el)
        else:
            # Get the python-docx Table object
            tbl_obj = None
            for t in doc.tables:
                if t._tbl is tbl_el:
                    tbl_obj = t
                    break
            if tbl_obj is None:
                continue

            # Get template data row (row index 1)
            if len(tbl_obj.rows) < 2:
                continue
            template_row_el = tbl_obj.rows[1]._tr

            # Clear and rebuild: remove all data rows first
            for row in tbl_obj.rows[1:]:
                tbl_el.remove(row._tr)

            # Add rows for each book entry
            for i, book in enumerate(rows):
                title  = book.get("title", "").strip()
                blang  = book.get("language", lang).strip() or lang
                genre  = book.get("genre", "").strip()
                fmt    = book.get("format", "").strip()
                royalty= book.get("royalty", "50").strip() or "50"

                new_row_el = deepcopy(template_row_el)
                tbl_el.append(new_row_el)

                # Get cells of new row
                new_cells = new_row_el.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc")
                num_cells = len(new_cells)

                if num_cells == 6:  # Sl.No, Title, Language, Genre, Format, Royalty%
                    vals = [str(i+1), title, blang, genre, fmt, royalty]
                elif num_cells == 5:  # Title, Language, Genre, Format, Royalty%
                    vals = [title, blang, genre, fmt, royalty]
                else:
                    vals = [title, blang, genre, fmt, royalty][:num_cells]

                for ci, (cell_el, val) in enumerate(zip(new_cells, vals)):
                    written = False
                    for p_el in cell_el.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"):
                        for r_el in p_el.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"):
                            for t_el in r_el.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                                if not written:
                                    t_el.text = val
                                    written = True
                                else:
                                    t_el.text = ""
                    if not written:
                        # No existing runs — create a new run with text
                        from lxml import etree as _etree
                        _ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                        for p_el in cell_el.findall(f"{{{_ns}}}p"):
                            r_el = _etree.SubElement(p_el, f"{{{_ns}}}r")
                            t_el = _etree.SubElement(r_el, f"{{{_ns}}}t")
                            t_el.text = val
                            break


# ── Fill Katha ─────────────────────────────────────────────────
def _set_para_title_text(para_el, title_name, para_text_orig, seq, ns):
    """Helper: write 'Title #N : <title_name>' into a paragraph element."""
    all_runs = para_el.findall(f".//{{{ns}}}r")
    # Build new label e.g. "Title #3 :"
    m = re.match(r"(Title\s*#\s*\d+\s*[:\-–—])", para_text_orig, re.IGNORECASE)
    sep = m.group(1)[-1] if m else ":"
    new_label = f"Title #{seq + 1} {sep} {title_name}"
    if all_runs:
        first_t = all_runs[0].find(f"{{{ns}}}t")
        if first_t is not None:
            first_t.text = new_label
        for r in all_runs[1:]:
            for t in r.findall(f"{{{ns}}}t"):
                t.text = ""
        first_run_ts = all_runs[0].findall(f"{{{ns}}}t")
        for t in first_run_ts[1:]:
            t.text = ""
    else:
        r_el = ET.SubElement(para_el, f"{{{ns}}}r")
        t_el = ET.SubElement(r_el, f"{{{ns}}}t")
        t_el.text = new_label


def fill_katha_annexure_a(doc, annexure_data, fields):
    """
    Fill Katha Annexure A.
    Finds 'Title #N :' paragraphs in the template and fills them.
    If user selected MORE titles than the template has blocks, clones the last
    title-paragraph + table block for each extra title and inserts them in order.
    annexure_data['katha_titles'] = [{title, language, genre, format, plan, addon, amount}, ...]
    """
    from copy import deepcopy
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    katha_titles = annexure_data.get("katha_titles", [])
    if not katha_titles:
        return

    lang_default   = fields.get("language", "Tamil").strip()
    plan_default   = fields.get("plan", "").strip()
    fmt_default    = fields.get("format", "").strip()
    amount_default = fields.get("amount", "").strip()

    body     = doc.element.body
    children = list(body)

    # ── Step 1: locate existing Title #N paragraph+table blocks in template ──
    title_blocks = []   # list of (para_idx, para_el, para_text, tbl_el or None)
    for i, child in enumerate(children):
        if child.tag.endswith("}p"):
            texts = [r.text for r in child.findall(f".//{{{ns}}}t") if r.text]
            full  = "".join(texts).strip()
            if re.match(r"Title\s*#\s*\d+\s*[:\-–—]", full, re.IGNORECASE):
                # find the next table sibling
                tbl_el = None
                for k in range(i + 1, min(i + 6, len(children))):
                    if children[k].tag.endswith("}tbl"):
                        tbl_el = children[k]
                        break
                title_blocks.append((i, child, full, tbl_el))

    if not title_blocks:
        return

    num_template = len(title_blocks)
    num_needed   = len(katha_titles)

    # ── Step 2: if we need MORE blocks than the template provides, clone & insert ──
    if num_needed > num_template:
        # Use the LAST template block as the clone source
        last_para_idx, last_para_el, last_para_text, last_tbl_el = title_blocks[-1]

        # Find insertion point: right after the last table (or last para if no table)
        if last_tbl_el is not None:
            insert_after_el = last_tbl_el
        else:
            insert_after_el = last_para_el

        # Re-read children after any previous insertions
        children = list(body)
        insert_pos = list(body).index(insert_after_el)

        for extra_seq in range(num_template, num_needed):
            # Clone paragraph
            new_para = deepcopy(last_para_el)
            # Update the Title #N label in the clone to the correct sequence number
            new_runs = new_para.findall(f".//{{{ns}}}r")
            sep_char = ":"
            m = re.match(r"Title\s*#\s*\d+\s*([:\-–—])", last_para_text, re.IGNORECASE)
            if m: sep_char = m.group(1)
            if new_runs:
                first_t = new_runs[0].find(f"{{{ns}}}t")
                if first_t is not None:
                    first_t.text = f"Title #{extra_seq + 1} {sep_char}"
                for r in new_runs[1:]:
                    for t in r.findall(f"{{{ns}}}t"):
                        t.text = ""

            # Clone table
            new_tbl = deepcopy(last_tbl_el) if last_tbl_el is not None else None

            # Clear data row in cloned table so old values don't bleed through
            if new_tbl is not None:
                tbl_rows = new_tbl.findall(f".//{{{ns}}}tr")
                if len(tbl_rows) >= 2:
                    for cell_el in tbl_rows[1].findall(f".//{{{ns}}}tc"):
                        for t_el in cell_el.findall(f".//{{{ns}}}t"):
                            t_el.text = ""

            # Insert: para then table, right after current insert_pos
            body.insert(insert_pos + 1, new_para)
            insert_pos += 1
            if new_tbl is not None:
                body.insert(insert_pos + 1, new_tbl)
                insert_pos += 1

        # Rebuild children list after insertions
        children = list(body)

        # Re-scan to pick up newly inserted blocks
        title_blocks = []
        for i, child in enumerate(children):
            if child.tag.endswith("}p"):
                texts = [r.text for r in child.findall(f".//{{{ns}}}t") if r.text]
                full  = "".join(texts).strip()
                if re.match(r"Title\s*#\s*\d+\s*[:\-–—]", full, re.IGNORECASE):
                    tbl_el = None
                    for k in range(i + 1, min(i + 6, len(children))):
                        if children[k].tag.endswith("}tbl"):
                            tbl_el = children[k]
                            break
                    title_blocks.append((i, child, full, tbl_el))

    # ── Step 3: fill each block with data ──
    for seq, (para_idx, para_el, para_text_orig, tbl_el) in enumerate(title_blocks):
        if seq >= num_needed:
            break

        title_data = katha_titles[seq]
        title_name = title_data.get("title", "").strip()
        lang       = title_data.get("language", lang_default).strip() or lang_default
        genre      = title_data.get("genre", "").strip()
        fmt        = title_data.get("format", fmt_default).strip() or fmt_default
        plan       = title_data.get("plan", plan_default).strip() or plan_default
        addon      = title_data.get("addon", "None").strip() or "None"
        amount     = title_data.get("amount", amount_default).strip() or amount_default

        if title_name:
            _set_para_title_text(para_el, title_name, para_text_orig, seq, ns)

        if tbl_el is None:
            continue

        # Get python-docx Table object for this tbl element
        tbl_obj = None
        for t in doc.tables:
            if t._tbl is tbl_el:
                tbl_obj = t
                break
        if tbl_obj is None or len(tbl_obj.rows) < 2:
            continue

        cells = tbl_obj.rows[1].cells
        vals  = [lang, genre, fmt, plan, addon, amount]
        for ci, val in enumerate(vals):
            if ci < len(cells):
                wc(cells[ci], val)


def fill_katha_doc(template_path, fields, director, annexure_formats, annexure_data=None):
    if annexure_data is None: annexure_data = {}
    doc = Document(template_path)
    a_name   = fields.get("author_name","").strip()
    pen      = fields.get("pen_name","").strip()
    pen_sfx  = f" (Pen Name: {pen})" if pen else ""
    parent   = fields.get("parent_name","").strip()
    addr     = fields.get("address","").strip()
    email    = fields.get("email","").strip()
    mobile   = fields.get("mobile","").strip()
    ref      = fields.get("ref_number","").strip()
    date_val = fields.get("date","").strip()
    lang     = fields.get("language","Tamil").strip()
    plan     = fields.get("plan","").strip()
    fmt      = fields.get("format","").strip()
    amount   = fields.get("amount","").strip()
    nom_name = fields.get("nominee_name","").strip()
    nom_rel  = fields.get("nominee_relationship","").strip()
    nom_addr = fields.get("nominee_address","").strip()
    nom_age  = fields.get("nominee_age","").strip()
    nom_me   = fields.get("nominee_mobile_email","").strip()
    nation   = fields.get("nationality","Indian").strip()

    # Determine effective format string from selected annexure formats
    eff_fmt = fmt if fmt else build_format_str(annexure_formats)

    for para in doc.paragraphs:
        merge_runs(para)
        if not para.runs: continue
        t = para.runs[0].text

        # ── Ref number ──────────────────────────────────────────
        if t.strip().startswith("Ref:") and re.search(r"KAT/TAM/", t):
            set_text(para, f"Ref: {ref}"); continue

        # ── Effective Date: paragraph is just '______ ("Effective Date")' (Silver++)
        #    or 'entered into on this _______ ("Effective Date")' (SGP template) ──
        if date_val and (
            ("Effective Date" in t and re.search(r"_{3,}", t)) or
            ("entered into on this" in t.lower() and re.search(r"_{3,}", t))
        ):
            nt = re.sub(r"_{3,}", date_val, t, count=1)
            set_text(para, nt); continue

        # ── Effective Date fallback: date already written as text ──
        if date_val and ("Effective Date" in t or "entered into on this" in t.lower()):
            m = re.search(r"(\d+(?:st|nd|rd|th)\s+\w+\s+\d{4})", t)
            if m:
                set_text(para, t.replace(m.group(1), date_val)); continue

        # ── Fill PACKAGE NAME in Annexure C ─────────────────────
        if "PACKAGE NAME" in t.upper() and re.search(r"_{3,}", t) and plan:
            nt = re.sub(r"_{3,}", plan.upper(), t, count=1)
            set_text(para, nt); continue

        # ── Author paragraph: '_____, son/daughter of _____, having address at ____.'
        #    Katha templates use underscores (not Mr./Ms./Mrs.) as placeholders ──
        if re.search(r"_{3,}", t) and re.search(r"son[/ ](?:daughter of|of)\s+", t, re.IGNORECASE):
            nt = t
            nt = re.sub(r"_{3,}", a_name + pen_sfx, nt, count=1)
            nt = re.sub(r"_{3,}", parent, nt, count=1)
            nt = re.sub(r"_{3,}", addr, nt, count=1)
            # In-line email / mobile if present in same paragraph (use lambda to avoid @ issues)
            _em, _mb2 = email, mobile
            nt = re.sub(r"(Email Id:\s*)_{3,}", lambda m, e=_em: m.group(1) + e, nt)
            nt = re.sub(r"(Mobile No:?[\s]*)_{3,}", lambda m, mb=_mb2: m.group(1) + mb, nt)
            set_text(para, nt); continue

        # ── Email / Mobile only paragraph (SGP template splits them to next para) ──
        if re.search(r"_{3,}", t) and "Email Id:" in t and "Mobile No" in t:
            nt = t
            _em, _mb2 = email, mobile
            nt = re.sub(r"(Email Id:\s*)_{3,}", lambda m, e=_em: m.group(1) + e, nt)
            nt = re.sub(r"(Mobile No:?[\s]*)_{3,}", lambda m, mb=_mb2: m.group(1) + mb, nt)
            set_text(para, nt); continue

    for t_idx, table in enumerate(doc.tables):
        rows = table.rows
        cols_n = len(table.columns)
        # Nominee table (table 0 - binding effect / nominee section)
        if t_idx == 0 and cols_n >= 4 and len(rows) >= 2:
            cells = rows[1].cells
            if len(cells) > 0: wc(cells[0], f"{nom_name}, {nom_age}, {nation}, {nom_addr}")
            if len(cells) > 1: wc(cells[1], nom_me.split("/")[0].strip() if "/" in nom_me else mobile)
            if len(cells) > 2: wc(cells[2], nom_me.split("/")[1].strip() if "/" in nom_me else email)
            if len(cells) > 3: wc(cells[3], nom_rel)
            if len(cells) > 4: wc(cells[4], "100")
        # Signature table (table 1)
        elif t_idx == 1:
            for row in rows:
                for cell in row.cells:
                    cell_text = cell.text
                    for para in cell.paragraphs:
                        merge_runs(para)
                        if not para.runs: continue
                        ct = para.runs[0].text
                        # Publisher signatory - always check first
                        if "Dr. Rajesh" in ct or "Prasanna" in ct:
                            nt = re.sub(r"Name:\s*(?:Dr\. Rajesh Devadas Ph\.D\.,|Mr\. Prasanna Devadas)", f"Name: {director}", ct)
                            if nt != ct: set_text(para, nt)
                        # Author name - only in "Signed For Author" cell
                        elif "Signed For Author" in cell_text or "Signed by Author" in cell_text:
                            if "Name:" in ct and a_name:
                                nt = re.sub(r"Name:\s*.*", f"Name: {a_name}{pen_sfx}", ct)
                                if nt != ct: set_text(para, nt)

    # Fill Katha Annexure A tables (Title #1, Title #2) with title name + row data
    fill_katha_annexure_a(doc, annexure_data, fields)

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

# ── Fill Pustaka ───────────────────────────────────────────────
def fill_pustaka_doc(template_path, fields, director, annexure_formats, annexure_data=None):
    if annexure_data is None: annexure_data = {}
    doc = Document(template_path)
    a_name   = fields.get("author_name","").strip()
    pen      = fields.get("pen_name","").strip()
    pen_sfx  = f" (Pen Name: {pen})" if pen else ""
    parent   = fields.get("parent_name","").strip()
    addr     = fields.get("address","").strip()
    email    = fields.get("email","").strip()
    mobile   = fields.get("mobile","").strip()
    ref      = fields.get("ref_number","").strip()
    date_val = fields.get("date","").strip()
    lang     = fields.get("language","Tamil").strip()
    nom_name = fields.get("nominee_name","").strip()
    nom_rel  = fields.get("nominee_relationship","").strip()
    nom_addr = fields.get("nominee_address","").strip()
    nom_age  = fields.get("nominee_age","").strip()
    nom_me   = fields.get("nominee_mobile_email","").strip()
    nation   = fields.get("nationality","Indian").strip()

    for para in doc.paragraphs:
        merge_runs(para)
        if not para.runs: continue
        t = para.runs[0].text
        if "Ref:" in t and "PUS/TAM/" in t:
            set_text(para, re.sub(r"PUS/TAM/\S*", ref, t)); continue
        m = re.search(r"(\d+(?:st|nd|rd|th)\s+\w+\s+\d{4})", t)
        if m and ("Effective Date" in t or "entered into on this" in t.lower()):
            set_text(para, t.replace(m.group(1), date_val)); continue
        if "son/daughter of" in t and re.search(r"_{3,}", t):
            nt = re.sub(r"_{3,}", a_name+pen_sfx, t, count=1)
            nt = re.sub(r"_{3,}", parent, nt, count=1)
            nt = re.sub(r"_{3,}", addr, nt, count=1)
            set_text(para, nt); continue
        if "Name:" in t and "Dr. Rajesh" not in t and "Witness" not in t and "Prasanna" not in t and a_name:
            nt = re.sub(r"Name:\s*\S[^\n]*", f"Name: {a_name}{pen_sfx}", t)
            if nt != t: set_text(para, nt); continue

    for table in doc.tables:
        cols_n = len(table.columns)
        rows = table.rows
        hdr0 = rows[0].cells[0].text.lower() if rows else ""
        if cols_n >= 4 and len(rows) >= 2 and "name" in hdr0 and ("age" in hdr0 or "nationality" in hdr0 or "address" in hdr0):
            cells = rows[1].cells
            if len(cells) > 0: wc(cells[0], f"{nom_name}, {nom_age}, {nation}, {nom_addr}")
            if len(cells) > 1: wc(cells[1], nom_me.split("/")[0].strip() if "/" in nom_me else mobile)
            if len(cells) > 2: wc(cells[2], nom_me.split("/")[1].strip() if "/" in nom_me else email)
            if len(cells) > 3: wc(cells[3], nom_rel)
            if len(cells) > 4: wc(cells[4], "100")
        for row in rows:
            for cell in row.cells:
                cell_text = cell.text
                for para in cell.paragraphs:
                    merge_runs(para)
                    if not para.runs: continue
                    ct = para.runs[0].text
                    # Publisher signatory - check first
                    if "Dr. Rajesh" in ct or "Prasanna" in ct:
                        nt = re.sub(r"Name:\s*(?:Dr\. Rajesh Devadas Ph\.D\.,|Mr\. Prasanna Devadas)", f"Name: {director}", ct)
                        if nt != ct: set_text(para, nt)
                    # Author name - only in "Signed For Author" cell
                    elif ("Signed For Author" in cell_text or "Signed by Author" in cell_text or
                          "Signed for and on behalf of the Author" in cell_text):
                        if "Name:" in ct and a_name:
                            nt = re.sub(r"Name:\s*.*", f"Name: {a_name}{pen_sfx}", ct)
                            if nt != ct: set_text(para, nt)

    # Fill Annexure A tables dynamically
    fill_annexure_a(doc, annexure_data, lang)

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

# ── Fill Addendum ──────────────────────────────────────────────
def fill_addendum_doc(template_path, fields, director, annexure_formats, annexure_data=None):
    if annexure_data is None: annexure_data = {}
    doc = Document(template_path)
    a_name    = fields.get("author_name","").strip()
    pen       = fields.get("pen_name","").strip()
    pen_sfx   = f" (Pen Name: {pen})" if pen else ""
    parent    = fields.get("parent_name","").strip()
    addr      = fields.get("address","").strip()
    email     = fields.get("email","").strip()
    mobile    = fields.get("mobile","").strip()
    ref       = fields.get("ref_number","").strip()
    add_date  = fields.get("addendum_date","").strip()
    orig_ref  = fields.get("orig_agreement_ref","").strip()
    orig_date = fields.get("orig_agreement_date","").strip()
    lang      = fields.get("language","Tamil").strip()
    plan      = fields.get("plan","").strip()
    amount    = fields.get("amount","").strip()
    fmt       = fields.get("format","").strip()
    eff_fmt   = fmt if fmt else build_format_str(annexure_formats)

    for para in doc.paragraphs:
        merge_runs(para)
        if not para.runs: continue
        t = para.runs[0].text

        # ── Ref number ──────────────────────────────────────────
        if t.strip().startswith("Ref:"):
            set_text(para, f"Ref: {ref}"); continue

        # ── Effective Date line: 'entered into on this ______ ("Effective Date").
        #    This is in reference to ... No. PUS/TAM/... dated ______'
        #    Template uses __________ blanks (not existing dates) ──
        if "Effective Date" in t and ("entered into on this" in t.lower() or "Addendum" in t):
            nt = t
            # Replace addendum date (first blank)
            if add_date:
                nt = re.sub(r"_{3,}", add_date, nt, count=1)
            # Replace original agreement ref number
            if orig_ref:
                nt = re.sub(r"[A-Z]{2,4}/TAM/[\w\-/]+", orig_ref, nt, count=1)
            # Replace original agreement date (second blank)
            if orig_date:
                nt = re.sub(r"_{3,}", orig_date, nt, count=1)
            set_text(para, nt); continue

        # ── PACKAGE NAME in Annexure C ───────────────────────────
        if "PACKAGE NAME" in t.upper() and re.search(r"_{3,}", t) and plan:
            nt = re.sub(r"_{3,}", plan.upper(), t, count=1)
            set_text(para, nt); continue

        # ── Author info line: '_____, son/daughter of _____, having address at _____
        #    (Email Id: _____ , Mobile : _____)' ──
        if re.search(r"son[/ ](?:daughter of|of)\s+", t, re.IGNORECASE) and re.search(r"_{3,}", t):
            nt = re.sub(r"_{3,}", a_name + pen_sfx, t, count=1)
            nt = re.sub(r"_{3,}", parent, nt, count=1)
            nt = re.sub(r"_{3,}", addr, nt, count=1)
            # Email and mobile also use __________ blanks in addendum template
            _em, _mb2 = email, mobile
            nt = re.sub(r"(Email Id\s*:\s*)_{3,}", lambda m, e=_em: m.group(1) + e, nt)
            nt = re.sub(r"(Mobile\s*:\s*)_{3,}", lambda m, mb=_mb2: m.group(1) + mb, nt)
            set_text(para, nt); continue

        # ── Author info fallback: Mr./Ms./Mrs. prefix style ──────
        if re.search(r"son[/ ](?:daughter of|of)\s+", t, re.IGNORECASE) and re.search(r"(Mr\.|Ms\.|Mrs\.)\s+\S", t):
            nt = re.sub(r"(Mr\.|Ms\.|Mrs\.)\s+[\w\.\s,]+?(?=,\s*son)", a_name+pen_sfx, t, count=1, flags=re.IGNORECASE)
            nt = re.sub(r"son[/ ](?:daughter of|of)\s+[\w\.\s]+?,", f"son/daughter of {parent},", nt, count=1, flags=re.IGNORECASE)
            nt = re.sub(r"having address at\s+[^,\(]+", f"having address at {addr}", nt, count=1)
            _em, _mb2 = email, mobile
            nt = re.sub(r"[\w\.\-]+@[\w\.\-]+", lambda m, e=_em: e, nt)
            nt = re.sub(r"\d{10}(?:\s*/\s*\d{10})?", lambda m, mb=_mb2: mb, nt)
            set_text(para, nt); continue

    for table in doc.tables:
        cols_n = len(table.columns)
        rows = table.rows
        # Signature table (2 cols, 1 row)
        if cols_n == 2 and len(rows) == 1:
            for cell in rows[0].cells:
                cell_text = cell.text
                for para in cell.paragraphs:
                    merge_runs(para)
                    if not para.runs: continue
                    ct = para.runs[0].text
                    # Publisher signatory - check first
                    if "Dr. Rajesh" in ct or "Prasanna" in ct:
                        nt = re.sub(r"Name:\s*(?:Dr\. Rajesh Devadas Ph\.D\.,|Mr\. Prasanna Devadas)", f"Name: {director}", ct)
                        if nt != ct: set_text(para, nt)
                    # Author name - only in "Signed For Author" cell
                    elif ("Signed For Author" in cell_text or "Signed by Author" in cell_text or
                          "Signed for and on behalf of the Author" in cell_text):
                        if "Name:" in ct and a_name:
                            nt = re.sub(r"Name:\s*.*", f"Name: {a_name}{pen_sfx}", ct)
                            if nt != ct: set_text(para, nt)
        # Katha plan/format/amount summary table (6 cols, language header)
        # This table is now filled via fill_katha_annexure_a below; skip here.
        pass

    # Fill Annexure A: katha addendums use Title #N paragraphs + 6-col table
    fill_katha_annexure_a(doc, annexure_data, fields)

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

# ── Master generate ────────────────────────────────────────────
def generate_document(fields, doc_type, director, annexure_formats,
                      pustaka_subtype=None, addendum_subtype=None,
                      addendum_plan=None, annexure_data=None):
    if doc_type == "katha_agreement":
        plan_l = fields.get("plan","").lower()
        key = "silver++" if any(x in plan_l for x in ["silver++","silver ++","sapphire","pearl"]) else "sgp"
        return fill_katha_doc(KATHA_TEMPLATES.get(key, KATHA_TEMPLATES["silver++"]), fields, director, annexure_formats, annexure_data or {})
    elif doc_type == "pustaka_agreement":
        tmpl = PUSTAKA_TEMPLATES.get(pustaka_subtype or "new_author", PUSTAKA_TEMPLATES["new_author"])
        return fill_pustaka_doc(tmpl, fields, director, annexure_formats, annexure_data or {})
    elif doc_type == "addendum":
        sub = addendum_subtype or "pustaka"
        if sub == "katha":
            tmpl = ADDENDUM_TEMPLATES["katha_pearl"] if "pearl" in (addendum_plan or "").lower() else ADDENDUM_TEMPLATES["katha_sapphire"]
        else:
            tmpl = ADDENDUM_TEMPLATES["pustaka"]
        return fill_addendum_doc(tmpl, fields, director, annexure_formats, annexure_data or {})
    raise ValueError(f"Unknown doc_type: {doc_type}")

# ── Text extraction ────────────────────────────────────────────
def extract_text(filepath):
    ext = filepath.rsplit(".", 1)[-1].lower()
    if ext == "txt":
        with open(filepath, encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif ext == "pdf":
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join(p.extract_text() or "" for p in reader.pages)
    elif ext in ("doc","docx"):
        doc = Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs)
    return ""

# ── AI helpers ─────────────────────────────────────────────────
def ai_extract_fields(raw_text, doc_type):
    fields_def = get_fields(doc_type)
    field_list = "\n".join(f"- {k}: {v[0]} ({v[1]})" for k, v in fields_def.items())
    prompt = f"""Extract author/agreement details from this text and return JSON.
Document type: {doc_type}
Fields to extract:
{field_list}

Rules:
- Return ONLY valid JSON with exact field keys above.
- If not found, use "".
- No markdown, no explanation.

Text:
\"\"\"{raw_text[:3000]}\"\"\"
JSON:"""
    try:
        client = anthropic.Anthropic(api_key=_load_api_key())
        resp = client.messages.create(model="claude-sonnet-4-6", max_tokens=1000,
                                      messages=[{"role":"user","content":prompt}])
        raw = resp.content[0].text.strip()
        raw = re.sub(r"^```json|^```|```$", "", raw, flags=re.MULTILINE).strip()
        return json.loads(raw)
    except Exception:
        return {}

def ai_extract_fields_from_image(image_path, doc_type):
    """Extract fields from a handwritten or scanned image using Claude's vision API."""
    import base64
    fields_def = get_fields(doc_type)
    field_list = "\n".join(f"- {k}: {v[0]} ({v[1]})" for k, v in fields_def.items())
    ext = image_path.rsplit(".", 1)[-1].lower()
    media_type_map = {
        "jpg": "image/jpeg", "jpeg": "image/jpeg",
        "png": "image/png", "webp": "image/webp",
        "bmp": "image/png",  # fallback
        "tiff": "image/png",
    }
    media_type = media_type_map.get(ext, "image/jpeg")
    with open(image_path, "rb") as f:
        image_data = base64.standard_b64encode(f.read()).decode("utf-8")
    prompt = f"""This image contains author/agreement details (possibly handwritten or printed).
Extract the following fields and return ONLY valid JSON with exact field keys below.
If a field is not found, use "".
Document type: {doc_type}
Fields to extract:
{field_list}

Return ONLY valid JSON. No markdown, no explanation."""
    try:
        client = anthropic.Anthropic(api_key=_load_api_key())
        resp = client.messages.create(
            model="claude-sonnet-4-6", max_tokens=1000,
            messages=[{"role":"user","content":[
                {"type":"image","source":{"type":"base64","media_type":media_type,"data":image_data}},
                {"type":"text","text":prompt}
            ]}]
        )
        raw = resp.content[0].text.strip()
        raw = re.sub(r"^```json|^```|```$", "", raw, flags=re.MULTILINE).strip()
        return json.loads(raw)
    except Exception:
        return {}

def ai_update_from_chat(user_msg, doc_type):
    field_keys = list(get_fields(doc_type).keys())
    prompt = f"""Extract any author/agreement details from this message.
Return JSON with only keys from: {field_keys}
If not found, omit the key. No markdown, no explanation.
Message: "{user_msg}"
JSON:"""
    try:
        client = anthropic.Anthropic(api_key=_load_api_key())
        resp = client.messages.create(model="claude-sonnet-4-6", max_tokens=256,
                                      messages=[{"role":"user","content":prompt}])
        raw = resp.content[0].text.strip()
        raw = re.sub(r"^```json|^```|```$", "", raw, flags=re.MULTILINE).strip()
        return json.loads(raw)
    except Exception:
        return {}

def ai_chat_reply(messages, doc_type, current_fields):
    fields_def = get_fields(doc_type)
    missing = [v[0] for k, v in fields_def.items() if v[2] and not current_fields.get(k,"").strip()]
    filled  = {v[0]: current_fields.get(k,"") for k, v in fields_def.items() if current_fields.get(k,"").strip()}
    system = f"""You are PustakaBot, a friendly AI assistant for Pustaka Digital Media Pvt. Ltd.
You help generate publishing agreements and addendums.
Current document type: {doc_type}
Fields already collected: {json.dumps(filled, ensure_ascii=False)}
Missing required fields: {missing}
Your job:
1. Extract and confirm author/agreement details from user messages.
2. Ask for missing required fields one at a time.
3. When all required fields are filled, say: "✅ All details collected! Ready to generate your document."
4. Keep responses concise and professional."""
    try:
        client = anthropic.Anthropic(api_key=_load_api_key())
        resp = client.messages.create(model="claude-sonnet-4-6", max_tokens=512,
                                      system=system, messages=messages[-10:])
        return resp.content[0].text.strip()
    except Exception as e:
        return f"⚠️ AI error: {e}"


# ══════════════════════════════════════════════════════════════
# MAIN APPLICATION WINDOW
# ══════════════════════════════════════════════════════════════
class PustakaApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Pustaka Agreement Assist")
        self.geometry("1200x780")
        self.minsize(960, 640)
        self.configure(fg_color=BG)
        # Set window icon (Windows: .ico, Mac: .icns)
        try:
            import platform
            if platform.system() == "Darwin":
                icon_path = _res("pustaka_icon.icns")
                if os.path.exists(icon_path):
                    self.iconphoto(True, tk.PhotoImage(file=icon_path))
            else:
                icon_path = _res("pustaka_icon.ico")
                if os.path.exists(icon_path):
                    self.iconbitmap(icon_path)
        except Exception:
            pass

        # State
        self.doc_type        = None
        self.pustaka_subtype = None
        self.addendum_subtype= None
        self.addendum_plan   = None
        self.fields          = {}
        self.chat_messages   = []
        self.director        = "Dr. Rajesh Devadas Ph.D.,"
        self.annexure_formats= []
        self.annexure_data   = {}
        self.generated_doc   = None
        self.doc_filename    = ""
        self.chat_history    = []

        self._build_layout()
        self._show_home()

    # ── Layout skeleton ────────────────────────────────────────
    def _build_layout(self):
        hdr = ctk.CTkFrame(self, fg_color=BLUE_DEEP, corner_radius=0, height=64)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        ctk.CTkLabel(hdr, text="🤖  Pustaka AI Agreement Generator",
                     font=("Segoe UI", 18, "bold"), text_color="white").pack(side="left", padx=24, pady=16)
        ctk.CTkLabel(hdr, text="Pustaka Digital Media Pvt. Ltd.",
                     font=FONT_SMALL, text_color="#93B4E0").pack(side="left")

        self.dir_var = ctk.StringVar(value=self.director)
        # ⚙️ API Key settings button
        ctk.CTkButton(hdr, text="⚙️ API Key", width=90, height=30,
                      fg_color=BLUE_MID, hover_color=BLUE_LIGHT,
                      font=FONT_SMALL, corner_radius=8,
                      command=self._show_api_settings).pack(side="right", padx=(0,12), pady=16)

        dir_frame = ctk.CTkFrame(hdr, fg_color="transparent")
        dir_frame.pack(side="right", padx=(20,4))
        ctk.CTkLabel(dir_frame, text="Signatory:", font=FONT_SMALL,
                     text_color="#93B4E0").pack(side="left", padx=(0,6))
        ctk.CTkOptionMenu(dir_frame,
                          values=["Dr. Rajesh Devadas Ph.D.,", "Mr. Prasanna Devadas"],
                          variable=self.dir_var, width=220,
                          fg_color=BLUE_MID, button_color=BLUE_LIGHT,
                          font=FONT_SMALL, command=self._on_director_change).pack(side="left")

        body = ctk.CTkFrame(self, fg_color=BG, corner_radius=0)
        body.pack(fill="both", expand=True)

        # Sidebar
        self.sidebar = ctk.CTkFrame(body, fg_color=WHITE, corner_radius=0,
                                    border_width=1, border_color=BORDER, width=200)
        self.sidebar.pack(side="left", fill="y")
        self.sidebar.pack_propagate(False)

        sidebar_hdr = ctk.CTkFrame(self.sidebar, fg_color=WHITE)
        sidebar_hdr.pack(fill="x", padx=12, pady=(14,4))
        ctk.CTkLabel(sidebar_hdr, text="📜 History",
                     font=FONT_HEAD, text_color=BLUE_DEEP).pack(side="left")
        ctk.CTkButton(sidebar_hdr, text="🗑 Clear", width=54, height=24,
                      fg_color=ERROR_RED, hover_color="#B91C1C",
                      font=("Segoe UI", 9), corner_radius=6,
                      command=self._clear_history).pack(side="right")

        ctk.CTkFrame(self.sidebar, fg_color=BORDER, height=1).pack(fill="x", padx=8)

        self.history_scroll = ctk.CTkScrollableFrame(self.sidebar, fg_color=WHITE, corner_radius=0)
        self.history_scroll.pack(fill="both", expand=True, padx=4, pady=4)

        self.content = ctk.CTkFrame(body, fg_color=BG, corner_radius=0)
        self.content.pack(side="left", fill="both", expand=True)

    def _on_director_change(self, val):
        self.director = val

    def _show_api_settings(self):
        """Popup to change API key."""
        win = ctk.CTkToplevel(self)
        win.title("API Key Settings")
        win.geometry("480x300")
        win.resizable(False, False)
        win.configure(fg_color=BG)
        win.grab_set()

        ctk.CTkLabel(win, text="🔑  Update API Key",
                     font=("Segoe UI", 15, "bold"), text_color=BLUE_DEEP).pack(pady=(24,4), padx=28, anchor="w")
        ctk.CTkLabel(win, text="Your current key is saved. Enter a new one to replace it.",
                     font=FONT_SMALL, text_color=TEXT_MUTED).pack(padx=28, anchor="w", pady=(0,16))

        key_frame = ctk.CTkFrame(win, fg_color=WHITE, corner_radius=10, border_width=1, border_color=BORDER)
        key_frame.pack(fill="x", padx=28, pady=(0,8))
        ctk.CTkLabel(key_frame, text="New API Key", font=FONT_SMALL, text_color=BLUE_DEEP).pack(anchor="w", padx=12, pady=(10,2))
        key_ent = ctk.CTkEntry(key_frame, placeholder_text="sk-ant-api03-...",
                                font=FONT_BODY, height=36, corner_radius=8,
                                border_color=BORDER, show="•")
        key_ent.pack(fill="x", padx=12, pady=(0,10))

        show_var = ctk.BooleanVar(value=False)
        def _toggle(): key_ent.configure(show="" if show_var.get() else "•")
        ctk.CTkCheckBox(win, text="Show key", variable=show_var, font=FONT_SMALL,
                        text_color=TEXT_MUTED, fg_color=BLUE_MID, hover_color=BLUE_LIGHT,
                        command=_toggle).pack(anchor="w", padx=28, pady=(0,12))

        status_lbl = ctk.CTkLabel(win, text="", font=FONT_SMALL, text_color=ERROR_RED)
        status_lbl.pack(padx=28, anchor="w")

        def _save():
            key = key_ent.get().strip()
            if not key:
                status_lbl.configure(text="⚠️  Please enter a key.", text_color=WARNING); return
            if not key.startswith("sk-ant-"):
                status_lbl.configure(text="⚠️  Key should start with sk-ant-...", text_color=WARNING); return
            _save_api_key(key)
            status_lbl.configure(text="✅  Saved!", text_color=SUCCESS)
            win.after(800, win.destroy)

        ctk.CTkButton(win, text="💾  Save Key", fg_color=SUCCESS, hover_color="#0A8060",
                      font=FONT_HEAD, height=40, corner_radius=10,
                      command=_save).pack(fill="x", padx=28, pady=(8,0))

    def _clear_content(self):
        for w in self.content.winfo_children():
            w.destroy()

    def _clear_history(self):
        if mb.askyesno("Clear History", "Are you sure you want to clear all chat history?"):
            self.chat_history = []
            self._refresh_history()

    def _refresh_history(self):
        for w in self.history_scroll.winfo_children():
            w.destroy()
        if not self.chat_history:
            ctk.CTkLabel(self.history_scroll, text="No history yet",
                         font=FONT_SMALL, text_color=TEXT_MUTED,
                         wraplength=170).pack(pady=10, padx=6)
            return
        for entry in self.chat_history:
            card = ctk.CTkFrame(self.history_scroll, fg_color=BLUE_PALE,
                                corner_radius=8, border_width=1, border_color=BORDER)
            card.pack(fill="x", pady=3, padx=2)
            ctk.CTkLabel(card, text=f"👤 {entry['author'][:18]}",
                         font=("Segoe UI", 10, "bold"), text_color=BLUE_DEEP,
                         anchor="w").pack(fill="x", padx=8, pady=(6,0))
            ctk.CTkLabel(card, text=entry['summary'],
                         font=FONT_SMALL, text_color=TEXT_MUTED,
                         anchor="w", wraplength=165).pack(fill="x", padx=8)
            ctk.CTkLabel(card, text=entry['timestamp'],
                         font=("Segoe UI", 9), text_color=TEXT_MUTED,
                         anchor="w").pack(fill="x", padx=8, pady=(0,6))

    def _save_to_history(self):
        dt_label = {"katha_agreement":"Katha","pustaka_agreement":"Pustaka","addendum":"Addendum"}.get(self.doc_type,"")
        plan = self.fields.get("plan","")
        summary = f"{dt_label}{' · '+plan if plan else ''}"
        entry = {
            "author":    self.fields.get("author_name","Unknown"),
            "summary":   summary,
            "filename":  self.doc_filename,
            "timestamp": datetime.now().strftime("%d %b %Y  %H:%M"),
        }
        self.chat_history.insert(0, entry)
        self.chat_history = self.chat_history[:20]
        self._refresh_history()

    # ══════════════════════════════════════════════════════════
    # SCREEN: HOME
    # ══════════════════════════════════════════════════════════
    def _show_home(self):
        self._clear_content()
        self.doc_type = None; self.fields = {}; self.chat_messages = []
        self.annexure_formats = []; self.annexure_data = {}
        self._refresh_history()

        wrapper = ctk.CTkFrame(self.content, fg_color=BG)
        wrapper.place(relx=0.5, rely=0.5, anchor="center")

        ctk.CTkLabel(wrapper, text="📄  What would you like to generate?",
                     font=FONT_TITLE, text_color=BLUE_DEEP).pack(pady=(0,32))

        cards = ctk.CTkFrame(wrapper, fg_color="transparent")
        cards.pack()

        self._home_card(cards, "📝", "Agreement",
                        "Katha or Pustaka\npublishing agreement",
                        lambda: self._show_agreement_type(), 0)
        self._home_card(cards, "📎", "Addendum",
                        "Amendment to an\nexisting agreement",
                        lambda: self._show_addendum_type(), 1)

    def _home_card(self, parent, icon, title, desc, cmd, col):
        card = ctk.CTkFrame(parent, fg_color=WHITE, corner_radius=16,
                            border_width=2, border_color=BORDER, width=260, height=200)
        card.grid(row=0, column=col, padx=16)
        card.pack_propagate(False)
        ctk.CTkLabel(card, text=icon, font=("Segoe UI", 40)).pack(pady=(28,8))
        ctk.CTkLabel(card, text=title, font=FONT_HEAD, text_color=BLUE_DEEP).pack()
        ctk.CTkLabel(card, text=desc, font=FONT_SMALL, text_color=TEXT_MUTED,
                     justify="center").pack(pady=6)
        ctk.CTkButton(card, text=f"Choose {title}", fg_color=BLUE_MID,
                      hover_color=BLUE_LIGHT, corner_radius=10, command=cmd).pack(pady=12)

    # ══════════════════════════════════════════════════════════
    # SCREEN: AGREEMENT TYPE
    # ══════════════════════════════════════════════════════════
    def _show_agreement_type(self):
        self._clear_content()
        wrap = ctk.CTkFrame(self.content, fg_color=BG)
        wrap.place(relx=0.5, rely=0.5, anchor="center")

        ctk.CTkLabel(wrap, text="📄  Which Agreement Type?",
                     font=FONT_TITLE, text_color=BLUE_DEEP).pack(pady=(0,28))
        row = ctk.CTkFrame(wrap, fg_color="transparent"); row.pack()

        options = [
            ("🎭", "Katha Agreement", "Fiction, poems, novels\nPearl / Sapphire / Silver++ / SGP", "katha_agreement"),
            ("📚", "Pustaka Agreement", "Broader publishing deal\nNew Author / Legal Heir / Licensor", None),
        ]
        for i,(icon,title,desc,dtype) in enumerate(options):
            cmd = (lambda d=dtype: self._start_flow(d)) if dtype else self._show_pustaka_subtype
            self._choice_card(row, icon, title, desc, cmd, i)

        ctk.CTkButton(wrap, text="← Back", fg_color="transparent",
                      text_color=BLUE_MID, hover_color=BLUE_PALE,
                      command=self._show_home).pack(pady=20)

    # ══════════════════════════════════════════════════════════
    # SCREEN: PUSTAKA SUBTYPE
    # ══════════════════════════════════════════════════════════
    def _show_pustaka_subtype(self):
        self._clear_content()
        wrap = ctk.CTkFrame(self.content, fg_color=BG)
        wrap.place(relx=0.5, rely=0.5, anchor="center")

        ctk.CTkLabel(wrap, text="📚  Pustaka Agreement — Author Type?",
                     font=FONT_TITLE, text_color=BLUE_DEEP).pack(pady=(0,28))
        row = ctk.CTkFrame(wrap, fg_color="transparent"); row.pack()

        opts = [
            ("✍️", "New Author",  "First-time publishing deal", "new_author"),
            ("⚖️", "Legal Heir",  "For deceased author's heir", "legal_heir"),
            ("📋", "Licensor",    "For rights licensor",        "licensor"),
        ]
        for i,(icon,title,desc,sub) in enumerate(opts):
            def cmd(s=sub): self.pustaka_subtype=s; self._start_flow("pustaka_agreement")
            self._choice_card(row, icon, title, desc, cmd, i)

        ctk.CTkButton(wrap, text="← Back", fg_color="transparent",
                      text_color=BLUE_MID, hover_color=BLUE_PALE,
                      command=self._show_agreement_type).pack(pady=20)

    # ══════════════════════════════════════════════════════════
    # SCREEN: ADDENDUM TYPE
    # ══════════════════════════════════════════════════════════
    def _show_addendum_type(self):
        self._clear_content()
        wrap = ctk.CTkFrame(self.content, fg_color=BG)
        wrap.place(relx=0.5, rely=0.5, anchor="center")

        ctk.CTkLabel(wrap, text="📎  Which Addendum Type?",
                     font=FONT_TITLE, text_color=BLUE_DEEP).pack(pady=(0,28))
        row = ctk.CTkFrame(wrap, fg_color="transparent"); row.pack()

        opts = [
            ("🎭","Katha Addendum", "Amendment for a Katha\nagreement (Pearl / Sapphire)", "katha"),
            ("📚","Pustaka Addendum","Amendment for a Pustaka\nagreement (Default)", "pustaka"),
        ]
        for i,(icon,title,desc,sub) in enumerate(opts):
            def cmd(s=sub):
                self.addendum_subtype = s
                if s == "katha": self._show_addendum_plan()
                else: self._start_flow("addendum")
            self._choice_card(row, icon, title, desc, cmd, i)

        ctk.CTkButton(wrap, text="← Back", fg_color="transparent",
                      text_color=BLUE_MID, hover_color=BLUE_PALE,
                      command=self._show_home).pack(pady=20)

    # ══════════════════════════════════════════════════════════
    # SCREEN: ADDENDUM PLAN
    # ══════════════════════════════════════════════════════════
    def _show_addendum_plan(self):
        self._clear_content()
        wrap = ctk.CTkFrame(self.content, fg_color=BG)
        wrap.place(relx=0.5, rely=0.5, anchor="center")

        ctk.CTkLabel(wrap, text="🎭  Katha Addendum — Plan?",
                     font=FONT_TITLE, text_color=BLUE_DEEP).pack(pady=(0,28))
        row = ctk.CTkFrame(wrap, fg_color="transparent"); row.pack()

        opts = [
            ("🔮","Pearl Plan","Pearl addendum template","pearl"),
            ("💎","Sapphire / Silver","Sapphire/Silver template","sapphire"),
        ]
        for i,(icon,title,desc,plan) in enumerate(opts):
            def cmd(p=plan): self.addendum_plan=p; self._start_flow("addendum")
            self._choice_card(row, icon, title, desc, cmd, i)

        ctk.CTkButton(wrap, text="← Back", fg_color="transparent",
                      text_color=BLUE_MID, hover_color=BLUE_PALE,
                      command=self._show_addendum_type).pack(pady=20)

    def _choice_card(self, parent, icon, title, desc, cmd, col):
        card = ctk.CTkFrame(parent, fg_color=WHITE, corner_radius=14,
                            border_width=2, border_color=BORDER, width=230, height=190)
        card.grid(row=0, column=col, padx=12)
        card.pack_propagate(False)
        ctk.CTkLabel(card, text=icon, font=("Segoe UI",32)).pack(pady=(22,4))
        ctk.CTkLabel(card, text=title, font=FONT_HEAD, text_color=BLUE_DEEP).pack()
        ctk.CTkLabel(card, text=desc, font=FONT_SMALL, text_color=TEXT_MUTED,
                     justify="center").pack(pady=4)
        ctk.CTkButton(card, text="Select", fg_color=BLUE_MID,
                      hover_color=BLUE_LIGHT, corner_radius=8,
                      height=32, command=cmd).pack(pady=8)

    # ══════════════════════════════════════════════════════════
    # MAIN FLOW SCREEN (Chat + Fields)
    # ══════════════════════════════════════════════════════════
    def _start_flow(self, doc_type):
        self.doc_type = doc_type
        self.fields   = {}
        self.chat_messages = []
        welcome = {
            "katha_agreement":   "Hi! 👋 I'll help generate a **Katha Agreement**.\n\nUpload a file with author details, or type them here.\n\nLet's start — what's the **author's full name**?",
            "pustaka_agreement": "Hi! 👋 I'll help generate a **Pustaka Agreement**.\n\nPlease provide the author's details. What's the **full name**?",
            "addendum":          "Hi! 👋 I'll help generate an **Addendum**.\n\nWhat's the **addendum reference number**? (e.g. KAT/TAM/325-04/2026)",
        }
        self.chat_messages.append({"role":"assistant","content": welcome.get(doc_type,"Hello! Let's get started.")})
        self._show_flow_screen()

    def _show_flow_screen(self):
        self._clear_content()
        doc_type   = self.doc_type
        fields_def = get_fields(doc_type)

        left  = ctk.CTkFrame(self.content, fg_color=BG, corner_radius=0)
        right = ctk.CTkFrame(self.content, fg_color=WHITE, corner_radius=0,
                             border_width=1, border_color=BORDER)
        left.pack(side="left", fill="both", expand=True, padx=(16,8), pady=16)
        right.pack(side="right", fill="y", padx=(0,16), pady=16, ipadx=8)
        right.configure(width=340); right.pack_propagate(False)

        # Progress bar
        prog_frame = ctk.CTkFrame(left, fg_color=WHITE, corner_radius=12,
                                  border_width=1, border_color=BORDER)
        prog_frame.pack(fill="x", pady=(0,10))
        total_req  = sum(1 for v in fields_def.values() if v[2])
        filled_req = sum(1 for k,v in fields_def.items() if v[2] and self.fields.get(k,"").strip())
        pct = filled_req / total_req if total_req else 0

        ctk.CTkLabel(prog_frame, text=f"Fields Collected: {filled_req}/{total_req} required",
                     font=FONT_SMALL, text_color=TEXT_MUTED).pack(anchor="w", padx=14, pady=(8,2))
        pb = ctk.CTkProgressBar(prog_frame, height=8, fg_color=BLUE_PALE,
                                 progress_color=BLUE_MID, corner_radius=4)
        pb.pack(fill="x", padx=14, pady=(0,10))
        pb.set(pct)

        # Chat window
        self.chat_frame = ctk.CTkScrollableFrame(left, fg_color=WHITE, corner_radius=12,
                                                  border_width=1, border_color=BORDER,
                                                  label_text="", height=340)
        self.chat_frame.pack(fill="both", expand=True, pady=(0,10))
        self._render_chat()

        # Input row
        inp_row = ctk.CTkFrame(left, fg_color="transparent")
        inp_row.pack(fill="x")
        self.chat_input = ctk.CTkEntry(inp_row, placeholder_text="Type author details or ask a question...",
                                        font=FONT_BODY, height=40, corner_radius=10,
                                        border_color=BORDER)
        self.chat_input.pack(side="left", fill="x", expand=True, padx=(0,8))
        self.chat_input.bind("<Return>", lambda e: self._on_send())
        ctk.CTkButton(inp_row, text="Send →", width=90, height=40,
                      fg_color=BLUE_MID, hover_color=BLUE_LIGHT,
                      corner_radius=10, command=self._on_send).pack(side="right")

        # Upload + Clear chat row
        upload_row = ctk.CTkFrame(left, fg_color="transparent")
        upload_row.pack(fill="x", pady=(10,0))
        ctk.CTkLabel(upload_row, text="Or upload a file with author details:",
                     font=FONT_SMALL, text_color=TEXT_MUTED).pack(side="left", padx=(0,10))
        ctk.CTkButton(upload_row, text="📁 Upload PDF / Word / TXT",
                      fg_color=WHITE, text_color=BLUE_MID,
                      border_width=2, border_color=BLUE_MID,
                      hover_color=BLUE_PALE, corner_radius=10,
                      command=self._on_upload).pack(side="left")
        ctk.CTkButton(upload_row, text="🗑 Clear Chat",
                      fg_color=WHITE, text_color=ERROR_RED,
                      border_width=2, border_color=ERROR_RED,
                      hover_color="#FEE2E2", corner_radius=10,
                      command=self._clear_chat).pack(side="left", padx=(8,0))

        # ── RIGHT: fields panel ────────────────────────────────
        hdr_row = ctk.CTkFrame(right, fg_color="transparent")
        hdr_row.pack(fill="x", padx=10, pady=(10,2))
        ctk.CTkLabel(hdr_row, text="📋  Collected Fields",
                     font=FONT_HEAD, text_color=BLUE_DEEP).pack(side="left")
        ctk.CTkButton(hdr_row, text="🗑 Clear Fields", width=100, height=26,
                      fg_color=WHITE, text_color=ERROR_RED,
                      border_width=1, border_color=ERROR_RED,
                      hover_color="#FEE2E2", corner_radius=8,
                      command=self._clear_fields).pack(side="right")

        # Katha: show which template will be used based on plan
        if doc_type == "katha_agreement":
            self._katha_tmpl_label = ctk.CTkLabel(right, text="📄 Template: Silver++ / Pearl / Sapphire template",
                                                   font=FONT_SMALL, text_color=BLUE_MID, anchor="w")
            self._katha_tmpl_label.pack(anchor="w", padx=14, pady=(0,4))
            self._update_katha_template_label()
        else:
            self._katha_tmpl_label = None

        fields_scroll = ctk.CTkScrollableFrame(right, fg_color=WHITE, corner_radius=0, height=400)
        fields_scroll.pack(fill="both", expand=True, padx=6)
        self.field_widgets = {}
        for k, v in fields_def.items():
            label, hint, required = v
            row_f = ctk.CTkFrame(fields_scroll, fg_color=WHITE, corner_radius=8,
                                 border_width=1, border_color=BORDER)
            row_f.pack(fill="x", pady=3, padx=2)
            icon = "🔴" if required and not self.fields.get(k,"").strip() else ("✅" if self.fields.get(k,"").strip() else "⚪")
            ctk.CTkLabel(row_f, text=f"{icon} {label}", font=FONT_SMALL,
                         text_color=BLUE_DEEP).pack(anchor="w", padx=8, pady=(4,0))
            entry = ctk.CTkEntry(row_f, placeholder_text=hint, font=FONT_SMALL,
                                 height=30, corner_radius=6, border_color=BORDER)
            entry.pack(fill="x", padx=8, pady=(2,6))
            val = self.fields.get(k,"")
            if val: entry.insert(0, val)
            entry.bind("<FocusOut>", lambda e, key=k, ent=entry: self._on_field_edit(key, ent))
            self.field_widgets[k] = entry

        # ── Annexure A ───────────────────────────────────────────
        if doc_type == "katha_agreement":
            # Katha: title-based UI (up to 2 titles)
            ctk.CTkFrame(fields_scroll, fg_color=BORDER, height=1).pack(fill="x", pady=(8,4), padx=2)
            ctk.CTkLabel(fields_scroll, text="📚 Annexure A — Title(s)",
                         font=("Segoe UI",10,"bold"), text_color=BLUE_DEEP).pack(anchor="w", padx=8, pady=(0,4))

            self.annexure_fmt_vars    = {}   # kept for compat
            self.annexure_rows_frames = {}
            self.annexure_num_vars    = {}

            existing_katha = self.annexure_data.get("katha_titles", [])
            num_var_k = ctk.IntVar(value=max(1, len(existing_katha)))
            self.annexure_num_vars["katha_titles"] = num_var_k

            katha_outer = ctk.CTkFrame(fields_scroll, fg_color=WHITE, corner_radius=8,
                                        border_width=1, border_color=BORDER)
            katha_outer.pack(fill="x", pady=2, padx=2)

            num_row_k = ctk.CTkFrame(katha_outer, fg_color="transparent")
            num_row_k.pack(fill="x", padx=8, pady=(6,2))
            ctk.CTkLabel(num_row_k, text="Number of titles (max 30):", font=FONT_SMALL,
                         text_color=TEXT_MUTED).pack(side="left")

            katha_rows_container = ctk.CTkFrame(katha_outer, fg_color="transparent")
            katha_rows_container.pack(fill="x", padx=8, pady=(0,8))

            def _rebuild_katha_rows(container=katha_rows_container):
                for w in container.winfo_children(): w.destroy()
                self.annexure_rows_frames["katha_titles"] = []
                n = min(30, self.annexure_num_vars["katha_titles"].get())
                existing = self.annexure_data.get("katha_titles", [])
                while len(existing) < n:
                    existing.append({"title":"","language":"","genre":"","format":"",
                                     "plan":"","addon":"None","amount":""})
                existing = existing[:n]
                katha_col_defs = [
                    ("title",    "Book Title",          6, 160),
                    ("language", "Language",            2,  70),
                    ("genre",    "Genre",               2,  70),
                    ("format",   "Format",              2,  70),
                    ("plan",     "Publishing Package",  3,  90),
                    ("addon",    "Add-on Services",     2,  80),
                    ("amount",   "Amount",              2,  70),
                ]
                for ri in range(n):
                    ctk.CTkLabel(container, text=f"Title #{ri+1}",
                                 font=FONT_SMALL, text_color=BLUE_DEEP).pack(anchor="w", pady=(6,0))
                    row_f = ctk.CTkFrame(container, fg_color="transparent")
                    row_f.pack(fill="x", pady=(0,2))
                    widgets = {}
                    for col_i, (col_key, ph, w_, minw) in enumerate(katha_col_defs):
                        ent = ctk.CTkEntry(row_f, placeholder_text=ph,
                                           font=FONT_SMALL, height=26,
                                           corner_radius=6, border_color=BORDER)
                        ent.grid(row=0, column=col_i, padx=(0,4), sticky="ew")
                        row_f.columnconfigure(col_i, weight=w_, minsize=minw)
                        val = existing[ri].get(col_key, "")
                        if not val and col_key == "addon": val = "None"
                        if val: ent.insert(0, val)
                        ent.bind("<FocusOut>", lambda e, ri2=ri, ck=col_key, en=ent:
                                 self._on_annexure_cell("katha_titles", ri2, ck, en))
                        widgets[col_key] = ent
                    self.annexure_rows_frames["katha_titles"].append(widgets)

            spin_fk = ctk.CTkFrame(num_row_k, fg_color="transparent")
            spin_fk.pack(side="left", padx=6)
            ctk.CTkButton(spin_fk, text="−", width=24, height=22,
                          fg_color=BORDER, text_color=BLUE_DEEP, hover_color=BLUE_PALE,
                          command=lambda: (
                              self.annexure_num_vars["katha_titles"].set(
                                  max(1, self.annexure_num_vars["katha_titles"].get()-1)),
                              _rebuild_katha_rows()
                          )).pack(side="left")
            ctk.CTkLabel(spin_fk, textvariable=num_var_k, font=FONT_SMALL,
                         text_color=BLUE_DEEP, width=24).pack(side="left")
            ctk.CTkButton(spin_fk, text="+", width=24, height=22,
                          fg_color=BORDER, text_color=BLUE_DEEP, hover_color=BLUE_PALE,
                          command=lambda: (
                              self.annexure_num_vars["katha_titles"].set(
                                  min(30, self.annexure_num_vars["katha_titles"].get()+1)),
                              _rebuild_katha_rows()
                          )).pack(side="left")

            _rebuild_katha_rows()

        elif doc_type in ("pustaka_agreement", "addendum"):
            # Katha addendum: use same Katha-style title-based Annexure A UI
            if doc_type == "addendum" and getattr(self, "addendum_subtype", None) == "katha":
                ctk.CTkFrame(fields_scroll, fg_color=BORDER, height=1).pack(fill="x", pady=(8,4), padx=2)
                ctk.CTkLabel(fields_scroll, text="📚 Annexure A — Title(s)",
                             font=("Segoe UI",10,"bold"), text_color=BLUE_DEEP).pack(anchor="w", padx=8, pady=(0,4))

                self.annexure_fmt_vars    = {}
                self.annexure_rows_frames = {}
                self.annexure_num_vars    = {}

                existing_katha = self.annexure_data.get("katha_titles", [])
                num_var_k = ctk.IntVar(value=max(1, len(existing_katha)))
                self.annexure_num_vars["katha_titles"] = num_var_k

                katha_outer = ctk.CTkFrame(fields_scroll, fg_color=WHITE, corner_radius=8,
                                            border_width=1, border_color=BORDER)
                katha_outer.pack(fill="x", pady=2, padx=2)

                num_row_k = ctk.CTkFrame(katha_outer, fg_color="transparent")
                num_row_k.pack(fill="x", padx=8, pady=(6,2))
                ctk.CTkLabel(num_row_k, text="Number of titles (max 30):", font=FONT_SMALL,
                             text_color=TEXT_MUTED).pack(side="left")

                katha_rows_container = ctk.CTkFrame(katha_outer, fg_color="transparent")
                katha_rows_container.pack(fill="x", padx=8, pady=(0,8))

                def _rebuild_katha_rows_add(container=katha_rows_container):
                    for w in container.winfo_children(): w.destroy()
                    self.annexure_rows_frames["katha_titles"] = []
                    n = min(30, self.annexure_num_vars["katha_titles"].get())
                    existing = self.annexure_data.get("katha_titles", [])
                    while len(existing) < n:
                        existing.append({"title":"","language":"","genre":"","format":"",
                                         "plan":"","addon":"None","amount":""})
                    existing = existing[:n]
                    katha_col_defs = [
                        ("title",    "Book Title",          6, 160),
                        ("language", "Language",            2,  70),
                        ("genre",    "Genre",               2,  70),
                        ("format",   "Format",              2,  70),
                        ("plan",     "Publishing Package",  3,  90),
                        ("addon",    "Add-on Services",     2,  80),
                        ("amount",   "Amount",              2,  70),
                    ]
                    for ri in range(n):
                        ctk.CTkLabel(container, text=f"Title #{ri+1}",
                                     font=FONT_SMALL, text_color=BLUE_DEEP).pack(anchor="w", pady=(6,0))
                        row_f = ctk.CTkFrame(container, fg_color="transparent")
                        row_f.pack(fill="x", pady=(0,2))
                        widgets = {}
                        for col_i, (col_key, ph, w_, minw) in enumerate(katha_col_defs):
                            ent = ctk.CTkEntry(row_f, placeholder_text=ph,
                                               font=FONT_SMALL, height=26,
                                               corner_radius=6, border_color=BORDER)
                            ent.grid(row=0, column=col_i, padx=(0,4), sticky="ew")
                            row_f.columnconfigure(col_i, weight=w_, minsize=minw)
                            val = existing[ri].get(col_key, "")
                            if not val and col_key == "addon": val = "None"
                            if val: ent.insert(0, val)
                            ent.bind("<FocusOut>", lambda e, ri2=ri, ck=col_key, en=ent:
                                     self._on_annexure_cell("katha_titles", ri2, ck, en))
                            widgets[col_key] = ent
                        self.annexure_rows_frames["katha_titles"].append(widgets)

                spin_fk2 = ctk.CTkFrame(num_row_k, fg_color="transparent")
                spin_fk2.pack(side="left", padx=6)
                ctk.CTkButton(spin_fk2, text="−", width=24, height=22,
                              fg_color=BORDER, text_color=BLUE_DEEP, hover_color=BLUE_PALE,
                              command=lambda: (
                                  self.annexure_num_vars["katha_titles"].set(
                                      max(1, self.annexure_num_vars["katha_titles"].get()-1)),
                                  _rebuild_katha_rows_add()
                              )).pack(side="left")
                ctk.CTkLabel(spin_fk2, textvariable=num_var_k, font=FONT_SMALL,
                             text_color=BLUE_DEEP, width=24).pack(side="left")
                ctk.CTkButton(spin_fk2, text="+", width=24, height=22,
                              fg_color=BORDER, text_color=BLUE_DEEP, hover_color=BLUE_PALE,
                              command=lambda: (
                                  self.annexure_num_vars["katha_titles"].set(
                                      min(30, self.annexure_num_vars["katha_titles"].get()+1)),
                                  _rebuild_katha_rows_add()
                              )).pack(side="left")

                _rebuild_katha_rows_add()

            else:
                # Pustaka / Pustaka Addendum: format-based UI (ebook / paperback / audiobook)
                ctk.CTkFrame(fields_scroll, fg_color=BORDER, height=1).pack(fill="x", pady=(8,4), padx=2)
                ctk.CTkLabel(fields_scroll, text="📚 Annexure A — Book Format(s)",
                             font=("Segoe UI",10,"bold"), text_color=BLUE_DEEP).pack(anchor="w", padx=8, pady=(0,4))

                self.annexure_fmt_vars    = {}
                self.annexure_rows_frames = {}
                self.annexure_num_vars    = {}
                fmt_options = [("ebook","📱 Ebook"), ("paperback","📖 Paperback"), ("audiobook","🎧 Audiobook")]

                for fmt_key, fmt_label in fmt_options:
                    var = ctk.BooleanVar(value=fmt_key in self.annexure_formats)
                    self.annexure_fmt_vars[fmt_key] = var

                    fmt_frame = ctk.CTkFrame(fields_scroll, fg_color=WHITE, corner_radius=8,
                                              border_width=1, border_color=BORDER)
                    fmt_frame.pack(fill="x", pady=2, padx=2)

                    ctk.CTkCheckBox(fmt_frame, text=fmt_label, variable=var,
                                    font=FONT_SMALL, text_color=BLUE_DEEP,
                                    fg_color=BLUE_MID, hover_color=BLUE_LIGHT,
                                    command=lambda fk=fmt_key: self._on_annexure_toggle(fk)).pack(anchor="w", padx=8, pady=(6,2))

                    num_row = ctk.CTkFrame(fmt_frame, fg_color="transparent")
                    num_row.pack(fill="x", padx=8, pady=(0,2))
                    ctk.CTkLabel(num_row, text="Number of titles:", font=FONT_SMALL,
                                 text_color=TEXT_MUTED).pack(side="left")
                    existing_rows = self.annexure_data.get(fmt_key, [])
                    num_var = ctk.IntVar(value=len(existing_rows))
                    self.annexure_num_vars[fmt_key] = num_var

                    rows_container = ctk.CTkFrame(fmt_frame, fg_color="transparent")
                    rows_container.pack(fill="x", padx=8, pady=(0,6))

                    def _rebuild_rows(fk=fmt_key, container=rows_container):
                        for w in container.winfo_children(): w.destroy()
                        self.annexure_rows_frames[fk] = []
                        n = self.annexure_num_vars[fk].get()
                        existing = self.annexure_data.get(fk, [])
                        while len(existing) < n:
                            existing.append({"title":"","language":"","genre":"","format":FMT_LABEL.get(fk,""),"royalty":"50"})
                        existing = existing[:n]
                        for ri in range(n):
                            ctk.CTkLabel(container, text=f"Title {ri+1}",
                                         font=FONT_SMALL, text_color=BLUE_DEEP).pack(anchor="w", pady=(4,0))
                            row_f = ctk.CTkFrame(container, fg_color="transparent")
                            row_f.pack(fill="x", pady=(0,2))
                            widgets = {}
                            col_defs = [
                                ("title",    "Book Title",  6, 160),
                                ("language", "Language",    2,  70),
                                ("genre",    "Genre",       2,  70),
                                ("format",   "Format",      2,  70),
                                ("royalty",  "%",           1,  36),
                            ]
                            for col_i, (col_key, ph, w_, minw) in enumerate(col_defs):
                                ent = ctk.CTkEntry(row_f, placeholder_text=ph,
                                                   font=FONT_SMALL, height=26,
                                                   corner_radius=6, border_color=BORDER)
                                ent.grid(row=0, column=col_i, padx=(0,4), sticky="ew")
                                row_f.columnconfigure(col_i, weight=w_, minsize=minw)
                                val = existing[ri].get(col_key, "")
                                if not val and col_key == "format":
                                    val = FMT_LABEL.get(fk,"")
                                if not val and col_key == "royalty":
                                    val = "50"
                                if val: ent.insert(0, val)
                                ent.bind("<FocusOut>", lambda e, fk2=fk, ri2=ri, ck=col_key, en=ent:
                                         self._on_annexure_cell(fk2, ri2, ck, en))
                                widgets[col_key] = ent
                            self.annexure_rows_frames[fk].append(widgets)

                    spin_f = ctk.CTkFrame(num_row, fg_color="transparent")
                    spin_f.pack(side="left", padx=6)
                    ctk.CTkButton(spin_f, text="−", width=24, height=22,
                                  fg_color=BORDER, text_color=BLUE_DEEP, hover_color=BLUE_PALE,
                                  command=lambda fk=fmt_key, rb=_rebuild_rows: (
                                      self.annexure_num_vars[fk].set(max(0, self.annexure_num_vars[fk].get()-1)),
                                      rb()
                                  )).pack(side="left")
                    ctk.CTkLabel(spin_f, textvariable=num_var, font=FONT_SMALL,
                                 text_color=BLUE_DEEP, width=24).pack(side="left")
                    ctk.CTkButton(spin_f, text="+", width=24, height=22,
                                  fg_color=BORDER, text_color=BLUE_DEEP, hover_color=BLUE_PALE,
                                  command=lambda fk=fmt_key, rb=_rebuild_rows: (
                                      self.annexure_num_vars[fk].set(min(20, self.annexure_num_vars[fk].get()+1)),
                                      rb()
                                  )).pack(side="left")

                    _rebuild_rows(fmt_key, rows_container)

        # Buttons divider
        ctk.CTkFrame(right, fg_color=BORDER, height=1).pack(fill="x", padx=10, pady=8)

        self.status_label = ctk.CTkLabel(right, text="", font=FONT_SMALL,
                                          text_color=TEXT_MUTED, wraplength=300)
        self.status_label.pack(padx=10, pady=(0,6))

        self.gen_btn = ctk.CTkButton(right, text="▶  Review & Generate",
                                     fg_color=SUCCESS, hover_color="#0A8060",
                                     font=FONT_HEAD, height=44, corner_radius=12,
                                     command=self._show_confirm)
        self.gen_btn.pack(fill="x", padx=10, pady=(0,4))

        ctk.CTkButton(right, text="← Back to Home", fg_color="transparent",
                      text_color=TEXT_MUTED, hover_color=BLUE_PALE,
                      command=self._show_home).pack(fill="x", padx=10)

        self._update_status()
        self._prog_bar = pb
        self._prog_total = total_req
        self._fields_scroll = fields_scroll

    def _render_chat(self):
        for w in self.chat_frame.winfo_children(): w.destroy()
        for msg in self.chat_messages:
            is_bot = msg["role"] == "assistant"
            row = ctk.CTkFrame(self.chat_frame, fg_color="transparent")
            row.pack(fill="x", pady=4, padx=4)
            bubble = ctk.CTkFrame(row,
                                   fg_color=BLUE_PALE if is_bot else BLUE_MID,
                                   corner_radius=12)
            text = ctk.CTkLabel(bubble, text=msg["content"].replace("**",""),
                                font=FONT_BODY,
                                text_color=BLUE_DEEP if is_bot else WHITE,
                                wraplength=420, justify="left")
            text.pack(padx=12, pady=8)
            if is_bot:
                bubble.pack(anchor="w", padx=(4,80))
            else:
                bubble.pack(anchor="e", padx=(80,4))
        self.chat_frame.after(50, lambda: self.chat_frame._parent_canvas.yview_moveto(1.0))

    def _on_field_edit(self, key, entry):
        self.fields[key] = entry.get().strip()
        self._update_status()
        # When the plan field changes for Katha, update the template indicator
        if key == "plan" and self.doc_type == "katha_agreement":
            self._update_katha_template_label()

    def _update_katha_template_label(self):
        """Show which Katha template will be used based on the selected plan."""
        if not hasattr(self, '_katha_tmpl_label'): return
        plan_l = self.fields.get("plan", "").lower()
        if any(x in plan_l for x in ["silver++", "silver ++", "sapphire", "pearl"]):
            tmpl_name = "Silver++ / Pearl / Sapphire template"
        elif any(x in plan_l for x in ["sgp", "silver gold", "silver", "gold"]):
            tmpl_name = "Silver Gold Plan template"
        else:
            tmpl_name = "Silver++ template (default)"
        self._katha_tmpl_label.configure(text=f"📄 Template: {tmpl_name}")

    def _update_status(self):
        if not hasattr(self, 'status_label'): return
        doc_type   = self.doc_type
        fields_def = get_fields(doc_type)
        missing    = [v[0] for k,v in fields_def.items() if v[2] and not self.fields.get(k,"").strip()]
        total_req  = sum(1 for v in fields_def.values() if v[2])
        filled_req = total_req - len(missing)
        if missing:
            self.status_label.configure(
                text=f"Still needed:\n• " + "\n• ".join(missing[:4]) + ("..." if len(missing)>4 else ""),
                text_color=WARNING)
            self.gen_btn.configure(state="normal")
        else:
            self.status_label.configure(text="✅ All required fields filled!", text_color=SUCCESS)
            self.gen_btn.configure(state="normal")
        if hasattr(self,'_prog_bar'):
            pct = filled_req / total_req if total_req else 0
            self._prog_bar.set(pct)

    def _on_send(self):
        msg = self.chat_input.get().strip()
        if not msg: return
        self.chat_input.delete(0, "end")
        # Sync manually-typed field widget values → self.fields BEFORE AI processes chat
        # This prevents manually entered data from being lost when AI syncs back
        for k, entry in self.field_widgets.items():
            val = entry.get().strip()
            if val: self.fields[k] = val
        self.chat_messages.append({"role":"user","content":msg})
        self._render_chat()

        def worker():
            extracted = ai_update_from_chat(msg, self.doc_type)
            for k,v in extracted.items():
                if v: self.fields[k] = v
            self.after(0, self._sync_fields_to_widgets)
            reply = ai_chat_reply(self.chat_messages, self.doc_type, self.fields)
            self.chat_messages.append({"role":"assistant","content":reply})
            self.after(0, self._render_chat)
            self.after(0, self._update_status)

        threading.Thread(target=worker, daemon=True).start()

    def _clear_chat(self):
        if mb.askyesno("Clear Chat", "Clear the conversation? Field data will be kept."):
            welcome = {
                "katha_agreement":   "Chat cleared. Fields are preserved — continue editing or click Review & Generate.",
                "pustaka_agreement": "Chat cleared. Fields are preserved — continue editing or click Review & Generate.",
                "addendum":          "Chat cleared. Fields are preserved — continue editing or click Review & Generate.",
            }
            self.chat_messages = [{"role":"assistant","content": welcome.get(self.doc_type,"Chat cleared.")}]
            self._render_chat()

    def _clear_fields(self):
        """Clear all collected field data and reset the widgets."""
        import tkinter.messagebox as _mb
        if not _mb.askyesno("Clear Fields", "Clear all collected field data? This cannot be undone."):
            return
        self.fields = {}
        self.annexure_data = {}
        self.annexure_formats = []
        # Reset all field entry widgets
        for k, entry in self.field_widgets.items():
            entry.delete(0, "end")
        # Reset annexure row widgets
        if hasattr(self, 'annexure_rows_frames'):
            for fmt_key, rows in self.annexure_rows_frames.items():
                for widgets in rows:
                    for col_key, ent in widgets.items():
                        ent.delete(0, "end")
                        if col_key == "addon":
                            ent.insert(0, "None")
                        elif col_key == "royalty":
                            ent.insert(0, "50")
        if hasattr(self, 'annexure_fmt_vars'):
            for var in self.annexure_fmt_vars.values():
                var.set(False)
        self._update_status()
        self._update_katha_template_label()

    def _sync_fields_to_widgets(self):
        for k, entry in self.field_widgets.items():
            val = self.fields.get(k,"")
            current = entry.get()
            if val and current != val:
                entry.delete(0,"end")
                entry.insert(0, val)

    def _on_annexure_toggle(self, fmt_key):
        if not hasattr(self, 'annexure_fmt_vars'): return
        self.annexure_formats = [k for k,v in self.annexure_fmt_vars.items() if v.get()]

    def _on_annexure_cell(self, fmt_key, row_index, col_key, entry):
        val = entry.get().strip()
        if fmt_key not in self.annexure_data:
            self.annexure_data[fmt_key] = []
        if fmt_key == "katha_titles":
            default_row = {"title":"","language":"","genre":"","format":"","plan":"","addon":"None","amount":""}
        else:
            default_row = {"title":"","language":"","genre":"","format":FMT_LABEL.get(fmt_key,""),"royalty":"50"}
        while len(self.annexure_data[fmt_key]) <= row_index:
            self.annexure_data[fmt_key].append(dict(default_row))
        self.annexure_data[fmt_key][row_index][col_key] = val

    def _collect_annexure_from_widgets(self):
        if not hasattr(self, 'annexure_rows_frames'): return
        if not self.annexure_rows_frames: return   # widgets already destroyed — keep existing data
        # Update annexure_formats only for non-katha (pustaka/addendum)
        if hasattr(self, 'annexure_fmt_vars') and self.annexure_fmt_vars:
            self.annexure_formats = [k for k,v in self.annexure_fmt_vars.items() if v.get()]
        new_data = {}
        for fmt_key, rows in self.annexure_rows_frames.items():
            collected = []
            for widgets in rows:
                try:
                    row_data = {ck: w.get().strip() for ck, w in widgets.items()}
                except Exception:
                    continue
                if fmt_key == "katha_titles":
                    if not row_data.get("addon"): row_data["addon"] = "None"
                else:
                    if not row_data.get("format"): row_data["format"] = FMT_LABEL.get(fmt_key,"")
                    if not row_data.get("royalty"): row_data["royalty"] = "50"
                collected.append(row_data)
            if collected:
                new_data[fmt_key] = collected
        self.annexure_data = new_data

    def _on_upload(self):
        path = fd.askopenfilename(
            title="Select file with author details",
            filetypes=[("Supported files","*.pdf *.docx *.doc *.txt *.png *.jpg *.jpeg *.bmp *.tiff *.webp"),
                       ("All files","*.*")]
        )
        if not path: return

        def worker():
            self.chat_messages.append({"role":"assistant","content":"📎 Reading your file... please wait."})
            self.after(0, self._render_chat)
            try:
                ext = path.rsplit(".", 1)[-1].lower()
                if ext in ("png","jpg","jpeg","bmp","tiff","webp"):
                    extracted = ai_extract_fields_from_image(path, self.doc_type)
                else:
                    raw = extract_text(path)
                    extracted = ai_extract_fields(raw, self.doc_type)
                for k,v in extracted.items():
                    if v: self.fields[k] = v
                fields_def = get_fields(self.doc_type)
                summary = "\n".join(f"• {fields_def[k][0]}: {v}"
                                    for k,v in extracted.items() if v and k in fields_def)
                if not summary: summary = "Couldn't extract details automatically. Please type them."
                reply = f"📎 Extracted from file:\n\n{summary}\n\nPlease review and correct if needed!"
            except Exception as e:
                reply = f"❌ File read error: {e}"
            self.chat_messages.append({"role":"assistant","content":reply})
            self.after(0, self._render_chat)
            self.after(0, self._sync_fields_to_widgets)
            self.after(0, self._update_status)

        threading.Thread(target=worker, daemon=True).start()

    # ══════════════════════════════════════════════════════════
    # SCREEN: CONFIRM & GENERATE
    # ══════════════════════════════════════════════════════════
    def _show_confirm(self):
        # Sync all field widget values → self.fields
        for k, entry in self.field_widgets.items():
            val = entry.get().strip()
            if val: self.fields[k] = val
        # Flush annexure widgets → self.annexure_data (must be done before widgets are destroyed)
        self._collect_annexure_from_widgets()
        self._clear_content()
        doc_type   = self.doc_type
        fields_def = get_fields(doc_type)

        ctk.CTkLabel(self.content, text="✅  Review & Confirm Details",
                     font=FONT_TITLE, text_color=BLUE_DEEP).pack(pady=20)

        scroll = ctk.CTkScrollableFrame(self.content, fg_color=WHITE, corner_radius=12,
                                         border_width=1, border_color=BORDER)
        scroll.pack(fill="both", expand=True, padx=24, pady=(0,16))

        self.confirm_entries = {}
        items = list(fields_def.items())
        grid = ctk.CTkFrame(scroll, fg_color=WHITE)
        grid.pack(fill="both", padx=16, pady=12)
        last_key = items[-1][0] if items else None
        for i, (k, v) in enumerate(items):
            label, hint, required = v
            col = i % 2
            row = i // 2
            cell = ctk.CTkFrame(grid, fg_color=WHITE)
            cell.grid(row=row, column=col, padx=8, pady=4, sticky="ew")
            grid.columnconfigure(col, weight=1)
            ctk.CTkLabel(cell, text=label+(" *" if required else ""),
                         font=FONT_SMALL, text_color=BLUE_DEEP).pack(anchor="w")
            ent = ctk.CTkEntry(cell, placeholder_text=hint, font=FONT_SMALL,
                               height=32, corner_radius=8, border_color=BORDER)
            ent.pack(fill="x")
            val = self.fields.get(k,"")
            if val: ent.insert(0, val)
            self.confirm_entries[k] = ent

        # ── Annexure A review ──────────────────────────────────
        if self.annexure_data:
            ctk.CTkFrame(scroll, fg_color=BORDER, height=1).pack(fill="x", padx=16, pady=(8,4))
            ctk.CTkLabel(scroll, text="📚 Annexure A — Selected Formats & Titles",
                         font=("Segoe UI",10,"bold"), text_color=BLUE_DEEP).pack(anchor="w", padx=16)
            for fmt_key, rows in self.annexure_data.items():
                if fmt_key == "katha_titles":
                    label_txt = f"  Titles — {len(rows)} title(s)"
                else:
                    label_txt = f"  {FMT_LABEL.get(fmt_key, fmt_key)} — {len(rows)} title(s)"
                ctk.CTkLabel(scroll, text=label_txt,
                             font=FONT_SMALL, text_color=BLUE_MID).pack(anchor="w", padx=24, pady=(4,0))
                for ri, row in enumerate(rows):
                    if fmt_key == "katha_titles":
                        detail = (f"    {ri+1}. {row.get('title','—')}  |  {row.get('language','—')}"
                                  f"  |  {row.get('genre','—')}  |  {row.get('format','—')}"
                                  f"  |  {row.get('plan','—')}  |  {row.get('addon','None')}"
                                  f"  |  {row.get('amount','—')}")
                    else:
                        detail = (f"    {ri+1}. {row.get('title','—')}  |  {row.get('language','—')}"
                                  f"  |  {row.get('genre','—')}  |  {row.get('format','—')}"
                                  f"  |  {row.get('royalty','50')}%")
                    ctk.CTkLabel(scroll, text=detail,
                                 font=FONT_SMALL, text_color=TEXT_MUTED).pack(anchor="w", padx=40, pady=1)

        # Signatory preview
        ctk.CTkFrame(scroll, fg_color=BORDER, height=1).pack(fill="x", padx=16, pady=(8,4))
        ctk.CTkLabel(scroll,
                     text=f"✍️  Signed for Author: {self.fields.get('author_name','—')}\n"
                          f"🏢  Signed for Publisher: {self.director}",
                     font=FONT_SMALL, text_color=BLUE_DEEP).pack(anchor="w", padx=16, pady=(0,8))

        # Download strip
        dl_strip = ctk.CTkFrame(self.content, fg_color=MINT, corner_radius=10,
                                border_width=1, border_color=SUCCESS)
        dl_strip.pack(fill="x", padx=24, pady=(0,8))
        ctk.CTkButton(dl_strip, text="⬇️  Download Word Document",
                      fg_color=SUCCESS, hover_color="#0A8060",
                      font=FONT_HEAD, height=40, corner_radius=8,
                      command=self._quick_save).pack(fill="x", padx=12, pady=8)

        # Bottom buttons
        btn_row = ctk.CTkFrame(self.content, fg_color=BG)
        btn_row.pack(fill="x", padx=24, pady=(0,16))
        ctk.CTkButton(btn_row, text="← Back to Edit", fg_color=WHITE,
                      text_color=BLUE_MID, border_width=2, border_color=BLUE_MID,
                      hover_color=BLUE_PALE, corner_radius=10,
                      command=self._show_flow_screen).pack(side="left")
        gen_btn = ctk.CTkButton(btn_row, text="🚀  Generate Document",
                      fg_color=BLUE_MID, hover_color=BLUE_LIGHT,
                      font=FONT_HEAD, height=44, corner_radius=12,
                      command=self._do_generate)
        gen_btn.pack(side="right")
        self._gen_btn_ref = gen_btn

    def _quick_save(self):
        """Collect confirm entries + generate + open save dialog."""
        for k, ent in self.confirm_entries.items():
            v = ent.get().strip()
            if v: self.fields[k] = v

        # annexure_data was already flushed in _show_confirm; don't re-collect (widgets gone)

        doc_type   = self.doc_type
        fields_def = get_fields(doc_type)
        missing = [v[0] for k,v in fields_def.items() if v[2] and not self.fields.get(k,"").strip()]
        if missing:
            mb.showerror("Missing Fields", "Still required:\n• " + "\n• ".join(missing))
            return

        try:
            doc_bytes = generate_document(
                fields=self.fields, doc_type=doc_type,
                director=self.director,
                annexure_formats=self.annexure_formats,
                pustaka_subtype=self.pustaka_subtype,
                addendum_subtype=self.addendum_subtype,
                addendum_plan=self.addendum_plan,
                annexure_data=self.annexure_data,
            )
            self.generated_doc = doc_bytes
            name = self.fields.get("author_name","author").replace(" ","_")
            date_str = datetime.now().strftime("%Y%m%d_%H%M")
            self.doc_filename = f"{doc_type}_{name}_{date_str}.docx"
            self._save_to_history()
            self._save_doc()
        except Exception as e:
            mb.showerror("Generation Error", str(e))

    def _do_generate(self):
        for k, ent in self.confirm_entries.items():
            v = ent.get().strip()
            if v: self.fields[k] = v

        doc_type   = self.doc_type
        fields_def = get_fields(doc_type)
        missing = [v[0] for k,v in fields_def.items() if v[2] and not self.fields.get(k,"").strip()]
        if missing:
            mb.showerror("Missing Fields", "Still required:\n• " + "\n• ".join(missing))
            return

        def worker():
            try:
                doc_bytes = generate_document(
                    fields=self.fields, doc_type=doc_type,
                    director=self.director,
                    annexure_formats=self.annexure_formats,
                    pustaka_subtype=self.pustaka_subtype,
                    addendum_subtype=self.addendum_subtype,
                    addendum_plan=self.addendum_plan,
                    annexure_data=self.annexure_data,
                )
                self.generated_doc = doc_bytes
                name = self.fields.get("author_name","author").replace(" ","_")
                date_str = datetime.now().strftime("%Y%m%d_%H%M")
                self.doc_filename = f"{doc_type}_{name}_{date_str}.docx"
                self.after(0, self._save_to_history)
                self.after(0, self._show_done)
            except Exception as e:
                self.after(0, lambda: mb.showerror("Generation Error", str(e)))

        try:
            self._gen_btn_ref.configure(text="⏳  Generating...", state="disabled")
        except Exception:
            pass
        threading.Thread(target=worker, daemon=True).start()

    # ══════════════════════════════════════════════════════════
    # SCREEN: DONE
    # ══════════════════════════════════════════════════════════
    def _show_done(self):
        self._clear_content()

        wrap = ctk.CTkFrame(self.content, fg_color=BG)
        wrap.place(relx=0.5, rely=0.5, anchor="center")
        self.after(300, self._save_doc)

        card = ctk.CTkFrame(wrap, fg_color=MINT, corner_radius=16,
                            border_width=2, border_color=SUCCESS)
        card.pack(pady=(0,24), padx=24)
        ctk.CTkLabel(card, text="🎉", font=("Segoe UI",48)).pack(pady=(24,4))
        ctk.CTkLabel(card, text="Your Document is Ready!",
                     font=("Segoe UI",20,"bold"), text_color="#0A6148").pack()
        ctk.CTkLabel(card, text=f"File: {self.doc_filename}",
                     font=FONT_SMALL, text_color=TEXT_MUTED).pack(pady=(4,20))

        btn_row = ctk.CTkFrame(wrap, fg_color="transparent")
        btn_row.pack()
        ctk.CTkButton(btn_row, text="⬇️  Save Word Document",
                      fg_color=SUCCESS, hover_color="#0A8060",
                      font=FONT_HEAD, height=48, width=220, corner_radius=12,
                      command=self._save_doc).grid(row=0, column=0, padx=8)
        ctk.CTkButton(btn_row, text="✏️  Edit Fields",
                      fg_color=WHITE, text_color=BLUE_MID,
                      border_width=2, border_color=BLUE_MID,
                      hover_color=BLUE_PALE, height=48, width=130, corner_radius=12,
                      command=self._show_confirm).grid(row=0, column=1, padx=8)
        ctk.CTkButton(btn_row, text="🆕  New Document",
                      fg_color=BLUE_MID, hover_color=BLUE_LIGHT,
                      height=48, width=150, corner_radius=12,
                      command=self._show_home).grid(row=0, column=2, padx=8)

        ctk.CTkLabel(wrap, text="📋  Document Summary",
                     font=FONT_HEAD, text_color=BLUE_DEEP).pack(pady=(28,8))
        fields_def = get_fields(self.doc_type)
        summary_frame = ctk.CTkFrame(wrap, fg_color=WHITE, corner_radius=12,
                                      border_width=1, border_color=BORDER)
        summary_frame.pack(padx=24, pady=(0,8))
        for k, v in list(fields_def.items())[:10]:
            val = self.fields.get(k,"") or "—"
            ctk.CTkLabel(summary_frame, text=f"{v[0]}: {val}",
                         font=FONT_SMALL, text_color=BLUE_DEEP).pack(anchor="w", padx=16, pady=2)
        ctk.CTkLabel(summary_frame,
                     text=f"Signed Author: {self.fields.get('author_name','—')}\nSignatory (Publisher): {self.director}",
                     font=FONT_SMALL, text_color=TEXT_MUTED).pack(anchor="w", padx=16, pady=(2,10))

    def _save_doc(self):
        if not self.generated_doc:
            mb.showerror("Error","No document to save.")
            return
        import tkinter as _tk
        import tkinter.filedialog as _fd
        _root = _tk.Tk()
        _root.withdraw()
        _root.attributes("-topmost", True)
        path = _fd.asksaveasfilename(
            parent=_root,
            title="Save Agreement Document",
            initialfile=self.doc_filename,
            defaultextension=".docx",
            filetypes=[("Word Document","*.docx"),("All files","*.*")]
        )
        _root.destroy()
        if path:
            with open(path,"wb") as f:
                f.write(self.generated_doc)
            mb.showinfo("Saved!", f"✅ Document saved to:\n{path}")


# ══════════════════════════════════════════════════════════════
# API KEY SETUP WINDOW
# ══════════════════════════════════════════════════════════════
class ApiKeySetup(ctk.CTk):
    """Shown on first launch (or when key is missing/invalid)."""
    def __init__(self, on_success):
        super().__init__()
        self.on_success = on_success
        self.title("Pustaka AI — Setup")
        self.geometry("520x420")
        self.resizable(False, False)
        self.configure(fg_color=BG)
        try:
            import platform as _pl
            if _pl.system() == "Darwin":
                icon_path = _res("pustaka_icon.icns")
                if os.path.exists(icon_path):
                    self.iconphoto(True, tk.PhotoImage(file=icon_path))
            else:
                icon_path = _res("pustaka_icon.ico")
                if os.path.exists(icon_path):
                    self.iconbitmap(icon_path)
        except Exception:
            pass
        self._build()

    def _build(self):
        # Header
        hdr = ctk.CTkFrame(self, fg_color=BLUE_DEEP, corner_radius=0, height=64)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        ctk.CTkLabel(hdr, text="🤖  Pustaka AI Agreement Generator",
                     font=("Segoe UI", 15, "bold"), text_color="white").pack(side="left", padx=24, pady=16)

        body = ctk.CTkFrame(self, fg_color=BG)
        body.pack(fill="both", expand=True, padx=32, pady=24)

        ctk.CTkLabel(body, text="🔑  Enter your Anthropic API Key",
                     font=("Segoe UI", 16, "bold"), text_color=BLUE_DEEP).pack(anchor="w", pady=(0, 8))

        ctk.CTkLabel(body,
                     text="This key is saved securely on your computer.\nYou only need to enter it once.",
                     font=FONT_SMALL, text_color=TEXT_MUTED, justify="left").pack(anchor="w", pady=(0, 16))

        # Key entry
        key_frame = ctk.CTkFrame(body, fg_color=WHITE, corner_radius=10,
                                  border_width=1, border_color=BORDER)
        key_frame.pack(fill="x", pady=(0, 8))
        ctk.CTkLabel(key_frame, text="API Key", font=FONT_SMALL,
                     text_color=BLUE_DEEP).pack(anchor="w", padx=12, pady=(10, 2))
        self.key_entry = ctk.CTkEntry(key_frame, placeholder_text="sk-ant-api03-...",
                                       font=FONT_BODY, height=38, corner_radius=8,
                                       border_color=BORDER, show="•")
        self.key_entry.pack(fill="x", padx=12, pady=(0, 10))

        # Show/hide toggle
        self.show_var = ctk.BooleanVar(value=False)
        ctk.CTkCheckBox(body, text="Show key", variable=self.show_var,
                        font=FONT_SMALL, text_color=TEXT_MUTED,
                        fg_color=BLUE_MID, hover_color=BLUE_LIGHT,
                        command=self._toggle_show).pack(anchor="w", pady=(0, 16))

        # Where to get key link info
        ctk.CTkLabel(body,
                     text="Get your key → console.anthropic.com → API Keys",
                     font=FONT_SMALL, text_color=BLUE_LIGHT).pack(anchor="w", pady=(0, 16))

        self.status = ctk.CTkLabel(body, text="", font=FONT_SMALL, text_color=ERROR_RED)
        self.status.pack(anchor="w", pady=(0, 4))

        ctk.CTkButton(body, text="✅  Save & Launch App",
                      fg_color=SUCCESS, hover_color="#0A8060",
                      font=FONT_HEAD, height=44, corner_radius=12,
                      command=self._on_save).pack(fill="x")

    def _toggle_show(self):
        self.key_entry.configure(show="" if self.show_var.get() else "•")

    def _on_save(self):
        key = self.key_entry.get().strip()
        if not key:
            self.status.configure(text="⚠️  Please enter your API key.")
            return
        if not key.startswith("sk-ant-"):
            self.status.configure(text="⚠️  Key should start with sk-ant-...")
            return
        self.status.configure(text="⏳  Verifying...", text_color=TEXT_MUTED)
        self.update()
        # Quick verify
        try:
            import anthropic as _anth
            c = _anth.Anthropic(api_key=key)
            c.messages.create(model="claude-haiku-4-5-20251001", max_tokens=10,
                              messages=[{"role":"user","content":"hi"}])
            _save_api_key(key)
            self.status.configure(text="✅  Key verified!", text_color=SUCCESS)
            self.after(600, lambda: (self.destroy(), self.on_success()))
        except Exception as e:
            err = str(e)
            if "authentication" in err.lower() or "api_key" in err.lower() or "401" in err:
                self.status.configure(text="❌  Invalid key. Check and try again.", text_color=ERROR_RED)
            else:
                # Network or other error — save anyway
                _save_api_key(key)
                self.after(400, lambda: (self.destroy(), self.on_success()))


# ══════════════════════════════════════════════════════════════
# ENTRY POINT
# ══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    key = _load_api_key()
    if not key:
        def launch_main():
            app = PustakaApp()
            app.mainloop()
        setup = ApiKeySetup(on_success=launch_main)
        setup.mainloop()
    else:
        app = PustakaApp()
        app.mainloop()
